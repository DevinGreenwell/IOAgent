"""Pytest configuration and fixtures for IOAgent tests."""

import os
import tempfile
import pytest
from pathlib import Path
from datetime import datetime
from typing import Generator, Dict, Any

from src.app_factory import create_app
from src.models.user import db, User
from src.models.project import Project
from src.models.evidence import Evidence
from src.models.timeline_entry import TimelineEntry
from src.models.causal_factor import CausalFactor


@pytest.fixture(scope='session')
def app():
    """Create application for testing."""
    # Create a temporary directory for test files
    with tempfile.TemporaryDirectory() as temp_dir:
        # Set test configuration
        os.environ['FLASK_ENV'] = 'testing'
        os.environ['TESTING'] = 'True'
        os.environ['DATABASE_URL'] = 'sqlite:///:memory:'
        os.environ['SECRET_KEY'] = 'test-secret-key'
        os.environ['JWT_SECRET_KEY'] = 'test-jwt-secret'
        os.environ['UPLOAD_FOLDER'] = temp_dir
        os.environ['PROJECTS_FOLDER'] = os.path.join(temp_dir, 'projects')
        
        # Create app with test config
        app = create_app('testing')
        
        # Set additional test config
        app.config['WTF_CSRF_ENABLED'] = False
        app.config['PRESERVE_CONTEXT_ON_EXCEPTION'] = False
        
        yield app


@pytest.fixture(scope='function')
def _db(app):
    """Create a clean database for each test."""
    with app.app_context():
        db.create_all()
        yield db
        db.session.remove()
        db.drop_all()


@pytest.fixture(scope='function')
def client(app, _db):
    """Create a test client."""
    return app.test_client()


@pytest.fixture(scope='function')
def runner(app):
    """Create a test CLI runner."""
    return app.test_cli_runner()


@pytest.fixture(scope='function')
def auth_headers(client, test_user):
    """Get authentication headers for requests."""
    # Login to get tokens
    response = client.post('/api/auth/login', json={
        'username': test_user.username,
        'password': 'TestPass123!'
    })
    
    assert response.status_code == 200
    tokens = response.get_json()
    
    return {
        'Authorization': f'Bearer {tokens["access_token"]}'
    }


@pytest.fixture(scope='function')
def admin_headers(client, admin_user):
    """Get admin authentication headers."""
    response = client.post('/api/auth/login', json={
        'username': admin_user.username,
        'password': 'AdminPass123!'
    })
    
    assert response.status_code == 200
    tokens = response.get_json()
    
    return {
        'Authorization': f'Bearer {tokens["access_token"]}'
    }


@pytest.fixture(scope='function')
def test_user(_db) -> User:
    """Create a test user."""
    user = User(
        username='testuser',
        email='test@example.com',
        role='user'
    )
    user.set_password('TestPass123!')
    
    _db.session.add(user)
    _db.session.commit()
    
    return user


@pytest.fixture(scope='function')
def admin_user(_db) -> User:
    """Create an admin user."""
    admin = User(
        username='admin',
        email='admin@example.com',
        role='admin'
    )
    admin.set_password('AdminPass123!')
    
    _db.session.add(admin)
    _db.session.commit()
    
    return admin


@pytest.fixture(scope='function')
def test_project(_db, test_user) -> Project:
    """Create a test project."""
    project = Project(
        name='Test Marine Incident',
        description='Test incident for unit testing',
        created_at=datetime.utcnow(),
        owner_id=test_user.id
    )
    
    _db.session.add(project)
    _db.session.commit()
    
    return project


@pytest.fixture(scope='function')
def test_evidence(_db, test_project) -> Evidence:
    """Create test evidence."""
    evidence = Evidence(
        project_id=test_project.id,
        title='Test Evidence Document',
        file_name='test_evidence.pdf',
        file_path='/test/path/test_evidence.pdf',
        file_type='application/pdf',
        file_size=1024,
        uploaded_at=datetime.utcnow(),
        summary='Test evidence summary'
    )
    
    _db.session.add(evidence)
    _db.session.commit()
    
    return evidence


@pytest.fixture(scope='function')
def test_timeline_entry(_db, test_project) -> TimelineEntry:
    """Create test timeline entry."""
    entry = TimelineEntry(
        project_id=test_project.id,
        timestamp=datetime.utcnow(),
        description='Test timeline event',
        event_type='incident',
        location='Test Location',
        actors='Test Actor',
        significance='high'
    )
    
    _db.session.add(entry)
    _db.session.commit()
    
    return entry


@pytest.fixture(scope='function')
def test_causal_factor(_db, test_project) -> CausalFactor:
    """Create test causal factor."""
    factor = CausalFactor(
        project_id=test_project.id,
        category='Human Factors',
        description='Test causal factor',
        barrier_type='Failed defense',
        remedial_action='Test remedial action',
        evidence_ids=[],
        contributing_factors=[]
    )
    
    _db.session.add(factor)
    _db.session.commit()
    
    return factor


@pytest.fixture(scope='function')
def mock_anthropic(monkeypatch):
    """Mock Anthropic API calls."""
    class MockMessage:
        def __init__(self, content):
            self.content = content
    
    class MockResponse:
        def __init__(self, content):
            self.content = [MockMessage(content)]
    
    class MockAnthropic:
        def __init__(self, api_key):
            self.messages = self
        
        def create(self, model, messages, **kwargs):
            # Return mock responses based on the prompt
            if any('timeline' in str(m) for m in messages):
                return MockResponse("Mock timeline suggestion")
            elif any('causal' in str(m) for m in messages):
                return MockResponse("Mock causal analysis")
            else:
                return MockResponse("Mock AI response")
    
    monkeypatch.setattr('anthropic.Anthropic', MockAnthropic)
    yield MockAnthropic


@pytest.fixture(scope='function')
def sample_pdf_file():
    """Create a sample PDF file for testing."""
    import io
    from reportlab.pdfgen import canvas
    
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer)
    p.drawString(100, 750, "Test PDF Document")
    p.drawString(100, 700, "This is a test document for IOAgent.")
    p.showPage()
    p.save()
    
    buffer.seek(0)
    return buffer


@pytest.fixture(scope='function')
def sample_docx_file():
    """Create a sample DOCX file for testing."""
    import io
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document for IOAgent.')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


@pytest.fixture(scope='function')
def api_test_data() -> Dict[str, Any]:
    """Provide common test data for API tests."""
    return {
        'project': {
            'name': 'API Test Project',
            'description': 'Project created via API'
        },
        'evidence': {
            'title': 'API Test Evidence',
            'summary': 'Evidence uploaded via API'
        },
        'timeline_entry': {
            'timestamp': '2024-01-15T10:30:00Z',
            'description': 'API test event',
            'event_type': 'incident',
            'location': 'Test Harbor',
            'actors': 'Test Vessel',
            'significance': 'high'
        },
        'causal_factor': {
            'category': 'Environmental',
            'description': 'Weather conditions',
            'barrier_type': 'Missing defense',
            'remedial_action': 'Implement weather monitoring'
        }
    }


@pytest.fixture(autouse=True)
def reset_database(_db):
    """Ensure database is clean before each test."""
    # This runs before each test due to autouse=True
    yield
    # Cleanup runs after each test
    _db.session.rollback()


# Pytest configuration
def pytest_configure(config):
    """Configure pytest with custom markers."""
    config.addinivalue_line(
        "markers", "slow: marks tests as slow (deselect with '-m \"not slow\"')"
    )
    config.addinivalue_line(
        "markers", "integration: marks tests as integration tests"
    )
    config.addinivalue_line(
        "markers", "unit: marks tests as unit tests"
    )
    config.addinivalue_line(
        "markers", "api: marks tests as API tests"
    )