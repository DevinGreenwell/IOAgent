"""Document generation service for creating reports."""

import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Service for generating various document formats."""
    
    def create_docx_report(self, content: Dict[str, Any], file_path: str, 
                          project: Any, include_appendices: bool = False) -> None:
        """
        Create a DOCX report from content.
        
        Args:
            content: Report content dictionary
            file_path: Output file path
            project: Project instance
            include_appendices: Whether to include appendices
        """
        try:
            doc = Document()
            
            # Set up styles
            self._setup_document_styles(doc)
            
            # Title page
            self._add_title_page(doc, project, content.get('title', 'Report of Investigation'))
            
            # Table of contents placeholder
            doc.add_page_break()
            doc.add_heading('Table of Contents', 1)
            doc.add_paragraph('[Table of Contents will be generated here]')
            
            # Executive Summary
            if 'executive_summary' in content:
                doc.add_page_break()
                doc.add_heading('Executive Summary', 1)
                doc.add_paragraph(content['executive_summary'])
            
            # Background
            if 'background' in content:
                doc.add_page_break()
                doc.add_heading('Background', 1)
                for paragraph in content['background'].split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Timeline of Events
            if 'timeline' in content:
                doc.add_page_break()
                doc.add_heading('Timeline of Events', 1)
                self._add_timeline_section(doc, content['timeline'])
            
            # Analysis
            if 'analysis' in content:
                doc.add_page_break()
                doc.add_heading('Analysis', 1)
                
                # Causal Factors
                if 'causal_factors' in content['analysis']:
                    doc.add_heading('Causal Factors', 2)
                    self._add_causal_factors_section(doc, content['analysis']['causal_factors'])
                
                # Contributing Factors
                if 'contributing_factors' in content['analysis']:
                    doc.add_heading('Contributing Factors', 2)
                    for factor in content['analysis']['contributing_factors']:
                        doc.add_paragraph(f"• {factor}", style='List Bullet')
            
            # Findings
            if 'findings' in content:
                doc.add_page_break()
                doc.add_heading('Findings', 1)
                for i, finding in enumerate(content['findings'], 1):
                    doc.add_paragraph(f"{i}. {finding}")
            
            # Recommendations
            if 'recommendations' in content:
                doc.add_page_break()
                doc.add_heading('Recommendations', 1)
                for i, rec in enumerate(content['recommendations'], 1):
                    doc.add_paragraph(f"{i}. {rec}")
            
            # Conclusions
            if 'conclusions' in content:
                doc.add_page_break()
                doc.add_heading('Conclusions', 1)
                doc.add_paragraph(content['conclusions'])
            
            # Appendices
            if include_appendices and 'appendices' in content:
                doc.add_page_break()
                doc.add_heading('Appendices', 1)
                self._add_appendices(doc, content['appendices'])
            
            # Save document
            doc.save(file_path)
            logger.info(f"DOCX report saved to {file_path}")
            
        except Exception as e:
            logger.error(f"Error creating DOCX report: {str(e)}")
            raise
    
    def create_pdf_report(self, content: Dict[str, Any], file_path: str, project: Any) -> None:
        """
        Create a PDF report from content.
        
        Args:
            content: Report content dictionary
            file_path: Output file path
            project: Project instance
        """
        try:
            # For now, create DOCX and note that PDF conversion would be needed
            # In production, you might use reportlab or convert the DOCX to PDF
            docx_path = file_path.replace('.pdf', '.docx')
            self.create_docx_report(content, docx_path, project)
            
            # TODO: Implement PDF conversion
            # This could use python-docx2pdf or similar library
            logger.warning(f"PDF generation not fully implemented. DOCX created at {docx_path}")
            
        except Exception as e:
            logger.error(f"Error creating PDF report: {str(e)}")
            raise
    
    def create_summary_document(self, summary_data: Dict[str, Any], file_path: str) -> None:
        """
        Create an executive summary document.
        
        Args:
            summary_data: Summary data dictionary
            file_path: Output file path
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Investigation Summary', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Project info
            doc.add_heading('Project Information', 1)
            doc.add_paragraph(f"Project: {summary_data['project_name']}")
            doc.add_paragraph(f"Description: {summary_data['description']}")
            doc.add_paragraph(f"Generated: {summary_data['generated_at']}")
            
            # Statistics
            doc.add_heading('Investigation Statistics', 1)
            stats = summary_data['statistics']
            doc.add_paragraph(f"• Evidence Files: {stats['evidence_files']}")
            doc.add_paragraph(f"• Timeline Events: {stats['timeline_events']}")
            doc.add_paragraph(f"• Causal Factors: {stats['causal_factors']}")
            
            # Key Findings
            if 'key_findings' in summary_data:
                doc.add_heading('Key Findings', 1)
                
                # Critical Events
                if summary_data['key_findings']['critical_events']:
                    doc.add_heading('Critical Events', 2)
                    for event in summary_data['key_findings']['critical_events']:
                        doc.add_paragraph(
                            f"• {event['timestamp']}: {event['description']}",
                            style='List Bullet'
                        )
                
                # Primary Causes
                if summary_data['key_findings']['primary_causes']:
                    doc.add_heading('Primary Causal Factors', 2)
                    for cause in summary_data['key_findings']['primary_causes']:
                        p = doc.add_paragraph()
                        p.add_run(f"{cause['category']}: ").bold = True
                        p.add_run(cause['description'])
                        if cause.get('remedial_action'):
                            doc.add_paragraph(
                                f"  Recommended Action: {cause['remedial_action']}",
                                style='List Bullet 2'
                            )
            
            # Save document
            doc.save(file_path)
            logger.info(f"Summary document saved to {file_path}")
            
        except Exception as e:
            logger.error(f"Error creating summary document: {str(e)}")
            raise
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Set up custom styles for the document."""
        # Modify existing styles
        styles = doc.styles
        
        # Heading 1
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        
        # Heading 2
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        
        # Normal text
        normal = styles['Normal']
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
    
    def _add_title_page(self, doc: Document, project: Any, title: str) -> None:
        """Add a title page to the document."""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Project name
        project_para = doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_run = project_para.add_run(project.name)
        project_run.font.size = Pt(18)
        
        # Vessel info
        if project.vessel_info:
            doc.add_paragraph()
            vessel_para = doc.add_paragraph()
            vessel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vessel_para.add_run(f"Vessel: {project.vessel_info.get('name', 'Unknown')}")
        
        # Incident info
        if project.incident_info:
            incident_para = doc.add_paragraph()
            incident_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            incident_para.add_run(f"Incident Date: {project.incident_info.get('date', 'Unknown')}")
            
            location_para = doc.add_paragraph()
            location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            location_para.add_run(f"Location: {project.incident_info.get('location', 'Unknown')}")
        
        # Generation date
        for _ in range(5):
            doc.add_paragraph()
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}")
    
    def _add_timeline_section(self, doc: Document, timeline_events: List[Dict[str, Any]]) -> None:
        """Add timeline section to the document."""
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Time'
        header_cells[1].text = 'Event'
        header_cells[2].text = 'Location'
        header_cells[3].text = 'Significance'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add events
        for event in timeline_events:
            row_cells = table.add_row().cells
            row_cells[0].text = event.get('timestamp', '')
            row_cells[1].text = event.get('description', '')
            row_cells[2].text = event.get('location', '')
            row_cells[3].text = event.get('significance', '')
    
    def _add_causal_factors_section(self, doc: Document, causal_factors: List[Dict[str, Any]]) -> None:
        """Add causal factors section to the document."""
        for factor in causal_factors:
            # Category as subheading
            category_para = doc.add_paragraph()
            category_para.add_run(factor.get('category', 'Unknown Category')).bold = True
            
            # Description
            doc.add_paragraph(factor.get('description', ''))
            
            # Barrier type
            if 'barrier_type' in factor:
                barrier_para = doc.add_paragraph()
                barrier_para.add_run('Barrier Type: ').bold = True
                barrier_para.add_run(factor['barrier_type'])
            
            # Remedial action
            if 'remedial_action' in factor:
                action_para = doc.add_paragraph()
                action_para.add_run('Recommended Action: ').bold = True
                action_para.add_run(factor['remedial_action'])
            
            # Add spacing
            doc.add_paragraph()
    
    def _add_appendices(self, doc: Document, appendices: List[Dict[str, Any]]) -> None:
        """Add appendices to the document."""
        for i, appendix in enumerate(appendices, 1):
            doc.add_heading(f"Appendix {i}: {appendix.get('title', 'Untitled')}", 2)
            
            if 'content' in appendix:
                for paragraph in appendix['content'].split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            if i < len(appendices):
                doc.add_page_break()