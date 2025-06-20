# ROI Document Generator - Core logic for generating USCG Reports of Investigation

import os
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json

from src.models.roi_models import InvestigationProject, ROIDocument, TimelineEntry, CausalFactor

class ROIGenerator:
    """Main class for generating ROI documents"""
    
    def __init__(self):
        self.project = None
        self.document = None
    
    def generate_roi(self, project: InvestigationProject, output_path: str) -> str:
        """Generate complete ROI document"""
        self.project = project
        self.document = Document()
        
        # Set up document formatting
        self._setup_document_formatting()
        
        # Generate each section
        self._generate_executive_summary()
        self._generate_preliminary_statement()
        self._generate_vessels_involved()
        self._generate_casualties_section()
        self._generate_findings_of_fact()
        self._generate_analysis()
        self._generate_conclusions()
        self._generate_actions_taken()
        self._generate_recommendations()
        
        # Save document
        self.document.save(output_path)
        return output_path
    
    def _setup_document_formatting(self):
        """Set up document styles and formatting"""
        # Set default font to Times New Roman 12pt
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Inches(12/72)  # 12pt
        
        # Create heading styles if they don't exist
        try:
            heading1 = self.document.styles['Heading 1']
        except:
            heading1 = self.document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        
        heading1.font.name = 'Times New Roman'
        heading1.font.size = Inches(14/72)  # 14pt
        heading1.font.bold = True
    
    def _generate_executive_summary(self):
        """Generate Executive Summary section"""
        # Title
        title = self._generate_title()
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary heading
        self.document.add_paragraph()
        summary_heading = self.document.add_paragraph()
        summary_run = summary_heading.add_run("EXECUTIVE SUMMARY")
        summary_run.bold = True
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        
        # Generate content if available
        if self.project.roi_document.executive_summary.scene_setting:
            self.document.add_paragraph(self.project.roi_document.executive_summary.scene_setting)
        
        if self.project.roi_document.executive_summary.outcomes:
            self.document.add_paragraph(self.project.roi_document.executive_summary.outcomes)
        
        if self.project.roi_document.executive_summary.causal_factors:
            self.document.add_paragraph(self.project.roi_document.executive_summary.causal_factors)
        
        # Page break
        self.document.add_page_break()
    
    def _generate_title(self) -> str:
        """Generate ROI title from project data"""
        vessels = []
        for vessel in self.project.vessels:
            if vessel.official_name and vessel.identification_number:
                vessels.append(f"{vessel.official_name} (O.N. {vessel.identification_number})")
        
        vessel_text = ", ".join(vessels) if vessels else "VESSEL"
        incident_type = self.project.incident_info.incident_type.upper() if self.project.incident_info.incident_type else "MARINE CASUALTY"
        
        # Check for casualties
        casualty_text = ""
        if self.project.incident_info.casualties_summary:
            casualty_text = " WITH INJURIES"
        
        location = self.project.incident_info.location.upper() if self.project.incident_info.location else "LOCATION"
        
        date_text = "DATE"
        if self.project.incident_info.incident_date:
            date_text = self.project.incident_info.incident_date.strftime("%B %d, %Y").upper()
        
        return f"{vessel_text}, {incident_type}{casualty_text} ON {location}, ON {date_text}"
    
    def _generate_preliminary_statement(self):
        """Generate Preliminary Statement section"""
        # Title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(self._generate_title())
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        
        # Section heading
        heading = self.document.add_paragraph()
        heading_run = heading.add_run("INVESTIGATING OFFICER'S REPORT")
        heading_run.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        
        # 1. Preliminary Statement
        self.document.add_paragraph("1. Preliminary Statement", style='Heading 1')
        
        # Required regulatory text
        self.document.add_paragraph(
            "This marine casualty investigation was conducted, and this report was "
            "submitted in accordance with Title 46, Code of Federal Regulations, "
            "Subpart 4.07, and under the authority of Title 46, United States Code, Chapter 63."
        )
        
        # Add custom preliminary statement if provided
        if self.project.roi_document.preliminary_statement:
            self.document.add_paragraph(self.project.roi_document.preliminary_statement)
    
    def _generate_vessels_involved(self):
        """Generate Vessels Involved section"""
        self.document.add_paragraph("2. Vessels Involved in the Incident", style='Heading 1')
        
        for vessel in self.project.vessels:
            # Create vessel information table
            table = self.document.add_table(rows=10, cols=2)
            table.style = 'Table Grid'
            
            # Populate table
            rows_data = [
                ("Official Name:", vessel.official_name or "N/A"),
                ("Identification Number:", vessel.identification_number or "N/A"),
                ("Flag:", vessel.flag or "N/A"),
                ("Vessel Class/Type/Sub-Type:", f"{vessel.vessel_class}/{vessel.vessel_type}/{vessel.vessel_subtype}" or "N/A"),
                ("Build Year:", str(vessel.build_year) if vessel.build_year else "N/A"),
                ("Gross Tonnage:", str(vessel.gross_tonnage) if vessel.gross_tonnage else "N/A"),
                ("Length:", str(vessel.length) if vessel.length else "N/A"),
                ("Beam/Width:", str(vessel.beam) if vessel.beam else "N/A"),
                ("Draft/Depth:", str(vessel.draft) if vessel.draft else "N/A"),
                ("Main/Primary Propulsion:", vessel.propulsion or "N/A")
            ]
            
            for i, (label, value) in enumerate(rows_data):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = value
            
            self.document.add_paragraph()
    
    def _generate_casualties_section(self):
        """Generate casualties section if applicable"""
        if self.project.personnel and any(p.status in ['injured', 'deceased'] for p in self.project.personnel):
            self.document.add_paragraph("3. Record of Deceased, Missing, and Injured", style='Heading 1')
            
            # Create casualties table
            table = self.document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Headers
            headers = ["Name/Role", "Age", "Status", "Description"]
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
                table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # Add casualty data
            for person in self.project.personnel:
                if person.status in ['injured', 'deceased']:
                    row = table.add_row()
                    row.cells[0].text = person.role
                    row.cells[1].text = "N/A"  # Age not typically included
                    row.cells[2].text = person.status.title()
                    row.cells[3].text = "See Findings of Fact"
            
            self.document.add_paragraph()
    
    def _generate_findings_of_fact(self):
        """Generate Findings of Fact section"""
        section_num = self._get_next_section_number()
        self.document.add_paragraph(f"{section_num}. Findings of Fact", style='Heading 1')
        
        if self.project.roi_document.findings_of_fact:
            for finding in self.project.roi_document.findings_of_fact:
                self.document.add_paragraph(finding.statement)
        else:
            # Generate from timeline if no explicit findings
            self._generate_findings_from_timeline()
    
    def _generate_findings_from_timeline(self):
        """Generate findings of fact from timeline entries"""
        if not self.project.timeline:
            self.document.add_paragraph("No findings of fact have been entered.")
            return
        
        # Sort timeline by timestamp
        sorted_timeline = sorted(
            [entry for entry in self.project.timeline if entry.timestamp],
            key=lambda x: x.timestamp
        )
        
        self.document.add_paragraph("The Incident:", style='Heading 2')
        
        for entry in sorted_timeline:
            # Format timestamp
            time_str = entry.timestamp.strftime("%B %d, %Y, at %H%M") if entry.timestamp else "At an unknown time"
            
            # Create finding statement
            finding_text = f"On {time_str}, {entry.description}"
            self.document.add_paragraph(finding_text)
    
    def _generate_analysis(self):
        """Generate Analysis section"""
        section_num = self._get_next_section_number()
        self.document.add_paragraph(f"{section_num}. Analysis", style='Heading 1')
        
        if self.project.roi_document.analysis_sections:
            for analysis in self.project.roi_document.analysis_sections:
                self.document.add_paragraph(analysis.title, style='Heading 2')
                self.document.add_paragraph(analysis.analysis_text)
        else:
            # Generate from causal factors
            self._generate_analysis_from_causal_factors()
    
    def _generate_analysis_from_causal_factors(self):
        """Generate analysis from causal factors"""
        if not self.project.causal_factors:
            self.document.add_paragraph("No causal factors have been identified.")
            return
        
        for factor in self.project.causal_factors:
            if factor.title:
                self.document.add_paragraph(factor.title, style='Heading 2')
            
            if factor.analysis_text:
                self.document.add_paragraph(factor.analysis_text)
            elif factor.description:
                self.document.add_paragraph(factor.description)
    
    def _generate_conclusions(self):
        """Generate Conclusions section"""
        section_num = self._get_next_section_number()
        self.document.add_paragraph(f"{section_num}. Conclusions", style='Heading 1')
        
        if self.project.roi_document.conclusions:
            for conclusion in self.project.roi_document.conclusions:
                self.document.add_paragraph(conclusion.statement)
        else:
            # Generate from causal factors
            self._generate_conclusions_from_causal_factors()
    
    def _generate_conclusions_from_causal_factors(self):
        """Generate conclusions from causal factors"""
        if not self.project.causal_factors:
            self.document.add_paragraph("No conclusions have been determined.")
            return
        
        # Find initiating event
        initiating_events = [entry for entry in self.project.timeline if entry.is_initiating_event]
        
        if initiating_events:
            initiating_event = initiating_events[0]
            conclusion_text = f"The Coast Guard determined the initiating event to be {initiating_event.description.lower()}."
            self.document.add_paragraph(conclusion_text)
        
        # List causal factors
        if self.project.causal_factors:
            factor_titles = [factor.title for factor in self.project.causal_factors if factor.title]
            if factor_titles:
                factors_text = f"Causal factors contributing to this casualty were: {', '.join(factor_titles)}."
                self.document.add_paragraph(factors_text)
    
    def _generate_actions_taken(self):
        """Generate Actions Taken section if applicable"""
        if self.project.roi_document.actions_taken:
            section_num = self._get_next_section_number()
            self.document.add_paragraph(f"{section_num}. Actions Taken Since the Incident", style='Heading 1')
            self.document.add_paragraph(self.project.roi_document.actions_taken)
    
    def _generate_recommendations(self):
        """Generate Recommendations section if applicable"""
        if self.project.roi_document.recommendations:
            section_num = self._get_next_section_number()
            self.document.add_paragraph(f"{section_num}. Recommendations", style='Heading 1')
            self.document.add_paragraph(self.project.roi_document.recommendations)
    
    def _get_next_section_number(self) -> int:
        """Get the next section number"""
        # Count existing sections (this is a simplified approach)
        # In a real implementation, you'd track section numbers more carefully
        sections_count = 2  # Start after Preliminary Statement and Vessels
        
        # Add 1 if casualties section exists
        if any(p.status in ['injured', 'deceased'] for p in self.project.personnel):
            sections_count += 1
        
        return sections_count + 1

class CausalAnalysisEngine:
    """Engine for causal analysis using USCG methodology"""
    
    def __init__(self):
        self.causal_categories = {
            'organization': 'Organization Factors',
            'workplace': 'Workplace Factors', 
            'precondition': 'Preconditions',
            'production': 'Production Factors',
            'defense': 'Defense Factors'
        }
    
    def analyze_timeline(self, timeline: List[TimelineEntry]) -> List[CausalFactor]:
        """Analyze timeline to identify potential causal factors"""
        causal_factors = []
        
        # Find events (adverse outcomes)
        events = [entry for entry in timeline if entry.type == 'event']
        
        for event in events:
            # Analyze each event for causal factors
            factors = self._analyze_event(event, timeline)
            causal_factors.extend(factors)
        
        return causal_factors
    
    def _analyze_event(self, event: TimelineEntry, timeline: List[TimelineEntry]) -> List[CausalFactor]:
        """Analyze a specific event for causal factors"""
        factors = []
        
        # Look for preceding actions and conditions
        event_time = event.timestamp
        if not event_time:
            return factors
        
        preceding_entries = [
            entry for entry in timeline 
            if entry.timestamp and entry.timestamp < event_time
        ]
        
        # Analyze for different factor categories
        for entry in preceding_entries:
            if entry.type == 'action':
                # Actions might indicate production or precondition factors
                factor = self._create_factor_from_action(entry, event)
                if factor:
                    factors.append(factor)
            elif entry.type == 'condition':
                # Conditions might indicate workplace or organization factors
                factor = self._create_factor_from_condition(entry, event)
                if factor:
                    factors.append(factor)
        
        return factors
    
    def _create_factor_from_action(self, action: TimelineEntry, event: TimelineEntry) -> CausalFactor:
        """Create causal factor from an action"""
        factor = CausalFactor()
        factor.event_id = event.id
        factor.category = 'production'  # Actions typically relate to production factors
        factor.title = f"Action leading to {event.description}"
        factor.description = action.description
        factor.evidence_support = action.evidence_ids
        return factor
    
    def _create_factor_from_condition(self, condition: TimelineEntry, event: TimelineEntry) -> CausalFactor:
        """Create causal factor from a condition"""
        factor = CausalFactor()
        factor.event_id = event.id
        factor.category = 'precondition'  # Conditions typically relate to preconditions
        factor.title = f"Condition contributing to {event.description}"
        factor.description = condition.description
        factor.evidence_support = condition.evidence_ids
        return factor

