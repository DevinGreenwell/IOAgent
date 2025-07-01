# USCG-Compliant ROI Document Generator
# Follows USCG Marine Investigation Documentation and Reporting Procedures Manual standards

import os
from datetime import datetime, date
from typing import List, Dict, Any, Optional, Union
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from zoneinfo import ZoneInfo  # Python 3.9+ standard tz database

from src.models.roi_models import InvestigationProject, ROIDocument, TimelineEntry, CausalFactor, Vessel, Personnel, Evidence

class USCGROIGenerator:
    """USCG-compliant ROI document generator following official standards"""
    
    def __init__(self) -> None:
        self.project: Optional[InvestigationProject] = None
        self.document: Optional[Document] = None
    
    def generate_roi(self, project: InvestigationProject, output_path: str) -> str:
        """Generate complete USCG-compliant ROI document"""
        self.project = project
        self.document = Document()
        
        # Set up USCG document formatting
        self._setup_uscg_formatting()
        
        # Generate USCG ROI sections in required order
        self._generate_executive_summary()
        self._generate_investigating_officers_report()
        
        # Save document
        self.document.save(output_path)
        return output_path
    
    def generate_roi_from_evidence_only(self, project: InvestigationProject, output_path: str) -> str:
        """Generate ROI directly from uploaded evidence files using AI - bypasses timeline/analysis workflow"""
        import logging
        logger = logging.getLogger('app')
        
        logger.info("🟡 DIRECT ROI: Starting AI-powered ROI generation from evidence files only")
        
        self.project = project
        self.document = Document()
        
        # Set up USCG document formatting
        self._setup_uscg_formatting()
        
        # Check if we have evidence to work with
        if not project.evidence_library or len(project.evidence_library) == 0:
            logger.error("🔴 DIRECT ROI: No evidence files available")
            raise ValueError("Cannot generate ROI: No evidence files uploaded")
        
        logger.info(f"🟡 DIRECT ROI: Processing {len(project.evidence_library)} evidence files")
        
        # Use AI to generate all content directly from evidence
        from src.models.anthropic_assistant import AnthropicAssistant
        ai_assistant = AnthropicAssistant()
        
        if not ai_assistant.client:
            logger.error("🔴 DIRECT ROI: No AI assistant available")
            raise ValueError("Cannot generate ROI: AI assistant not configured")
        
        # Generate comprehensive ROI sections using AI
        roi_content = self._generate_complete_roi_from_evidence(ai_assistant)
        
        if not roi_content:
            logger.error("🔴 DIRECT ROI: Failed to generate ROI content from evidence")
            raise ValueError("Failed to extract sufficient information from evidence files")
        
        # Generate document sections
        self._generate_ai_executive_summary(roi_content)
        self._generate_ai_investigating_officers_report(roi_content)
        
        # Save document
        self.document.save(output_path)
        logger.info(f"🟢 DIRECT ROI: Successfully generated ROI document at {output_path}")
        return output_path
    
    def _setup_uscg_formatting(self) -> None:
        """Set up USCG-required document formatting"""
        if not self.document:
            return
            
        # Set default font to Times New Roman 12-pitch as required
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set paragraph spacing
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.0
        
        # Set margins to 1 inch
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _format_date(self, date_obj: Optional[datetime]) -> str:
        """Format date according to USCG standard: Month DD, YYYY"""
        if date_obj:
            return date_obj.strftime("%B %d, %Y")
        return "[Date to be determined]"
    
    def _format_time(self, time_obj: Optional[datetime]) -> str:
        """Format time according to USCG standard (HHMM) and the incident's local time zone if provided"""
        if not time_obj:
            return "[Time unknown]"
        if not self.project:
            return time_obj.strftime("%H%M")
            
        tz_name = getattr(self.project.incident_info, "time_zone", None)
        if tz_name:
            try:
                time_obj = time_obj.astimezone(ZoneInfo(tz_name))
            except Exception:
                # Fallback silently if tz string is invalid
                pass
        return time_obj.strftime("%H%M")
    
    def _italicize_vessel_names(self, text: str) -> str:
        """Helper to mark vessel names for italicization"""
        # This is a placeholder - actual italicization happens when adding to document
        return text
    
    def _generate_executive_summary(self) -> None:
        """Generate Executive Summary section per USCG standards"""
        if not self.document or not self.project:
            return
            
        # Generate title in USCG format
        title = self._generate_uscg_title()
        
        # Title paragraph
        title_para = self.document.add_paragraph()
        title_para.add_run(title).bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        self.document.add_paragraph()
        
        # Executive Summary heading
        summary_heading = self.document.add_paragraph()
        summary_heading.add_run("EXECUTIVE SUMMARY").bold = True
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        
        # Use AI-generated executive summary if available
        if hasattr(self.project.roi_document, 'executive_summary') and self.project.roi_document.executive_summary:
            exec_summary = self.project.roi_document.executive_summary
            
            # Check if we have properly generated AI content
            if hasattr(exec_summary, 'scene_setting') and exec_summary.scene_setting:
                # Use the AI-generated comprehensive paragraphs
                self.document.add_paragraph(exec_summary.scene_setting)
                self.document.add_paragraph()  # Space between paragraphs
                
                if hasattr(exec_summary, 'outcomes') and exec_summary.outcomes:
                    self.document.add_paragraph(exec_summary.outcomes)
                    self.document.add_paragraph()  # Space between paragraphs
                
                if hasattr(exec_summary, 'causal_factors') and exec_summary.causal_factors:
                    self.document.add_paragraph(exec_summary.causal_factors)
            else:
                # Fallback to old methods if AI content is not available
                self._generate_fallback_executive_summary()
        else:
            # Fallback to old methods if no executive summary
            self._generate_fallback_executive_summary()
        
        self.document.add_page_break()
    
    def _generate_fallback_executive_summary(self) -> None:
        """Generate fallback executive summary using old methods"""
        # Paragraph 1: Scene Setting
        scene_para = self._generate_scene_setting_paragraph()
        self.document.add_paragraph(scene_para)
        
        self.document.add_paragraph()  # Space between paragraphs
        
        # Paragraph 2: Outcomes
        outcomes_para = self._generate_outcomes_paragraph()
        self.document.add_paragraph(outcomes_para)
        
        self.document.add_paragraph()  # Space between paragraphs
        
        # Paragraph 3: Causal Factors
        causal_para = self._generate_causal_factors_paragraph()
        self.document.add_paragraph(causal_para)
    
    def _generate_uscg_title(self) -> str:
        """Generate title in USCG format using Investigation Title and Official Number"""
        if not self.project:
            return "VESSEL INCIDENT"
            
        # Use Investigation Title and Official Number from project info
        investigation_title = self.project.metadata.title or "VESSEL INCIDENT"
        official_number = getattr(self.project, 'official_number', None)
        
        # Format official number
        if official_number and official_number != 'N/A':
            on_text = f"(O.N. {official_number})"
        else:
            on_text = ""
        
        # Build title in USCG format
        vessel_text = f"{investigation_title.upper()} {on_text}".strip()
        
        # Determine casualty type with injuries/fatalities
        casualty_type = self.project.incident_info.incident_type or "MARINE CASUALTY"
        
        # Check for personnel casualties
        has_injuries = False
        has_fatalities = False
        
        for entry in self.project.timeline:
            desc_lower = entry.description.lower()
            if any(word in desc_lower for word in ['injured', 'injury', 'injuries']):
                has_injuries = True
            if any(word in desc_lower for word in ['deceased', 'death', 'fatality', 'died']):
                has_fatalities = True
        
        if has_fatalities:
            casualty_desc = f"{casualty_type.upper()} WITH LOSS OF LIFE"
        elif has_injuries:
            casualty_desc = f"{casualty_type.upper()} WITH INJURIES"
        else:
            casualty_desc = casualty_type.upper()
        
        # Location
        location = self.project.incident_info.location or "LOCATION"
        if "near" not in location.lower() and "on" not in location.lower():
            location = f"ON {location.upper()}"
        else:
            location = location.upper()
        
        # Date
        date_str = self._format_date(self.project.incident_info.incident_date).upper()
        
        return f"{vessel_text}, {casualty_desc} {location} ON {date_str}"
    
    def _generate_scene_setting_paragraph(self) -> str:
        """Generate paragraph 1 of Executive Summary - Scene Setting"""
        if not self.project:
            return "Scene setting information not available."
            
        # Get incident date and time
        incident_date = self._format_date(self.project.incident_info.incident_date)
        
        # Find the earliest timeline entry for departure/start
        earliest_entry = None
        if self.project.timeline:
            sorted_timeline = sorted(self.project.timeline, key=lambda x: x.timestamp or datetime.max)
            earliest_entry = sorted_timeline[0] if sorted_timeline else None
        
        # Build scene setting
        if earliest_entry and earliest_entry.timestamp:
            time_str = self._format_time(earliest_entry.timestamp)
            opening = f"On {incident_date}, at {time_str}, "
        else:
            opening = f"On {incident_date}, "
        
        # Get vessel name
        vessel_name = ""
        if self.project.vessels and self.project.vessels[0].official_name:
            vessel_name = f"the {self.project.vessels[0].official_name} "
        
        # Extract key scene-setting details from timeline
        scene_details = []
        for entry in self.project.timeline[:5]:  # Look at first few entries
            if any(word in entry.description.lower() for word in ['departed', 'began', 'started', 'commenced']):
                scene_details.append(entry.description)
                break
        
        if scene_details:
            scene = opening + vessel_name + scene_details[0].lower()
        else:
            scene = opening + vessel_name + f"was operating in {self.project.incident_info.location}"
        
        # Add what happened
        incident_type = self.project.incident_info.incident_type or "an incident"
        scene += f". {self._describe_incident_from_timeline()}"
        
        return scene
    
    def _generate_outcomes_paragraph(self) -> str:
        """Generate paragraph 2 of Executive Summary - Outcomes"""
        if not self.project:
            return "Outcome information not available."
            
        outcomes = []
        
        # Scan timeline for outcomes
        for entry in self.project.timeline:
            desc_lower = entry.description.lower()
            
            # Personnel casualties
            if 'deceased' in desc_lower or 'died' in desc_lower or 'death' in desc_lower:
                if 'pronounced' in desc_lower:
                    outcomes.append(entry.description)
                else:
                    outcomes.append("resulted in loss of life")
            elif 'injured' in desc_lower or 'injury' in desc_lower:
                outcomes.append("resulted in personnel injuries")
            
            # Vessel damage
            elif any(word in desc_lower for word in ['damage', 'holed', 'flooded', 'fire', 'explosion']):
                outcomes.append("resulted in vessel damage")
            
            # Environmental
            elif any(word in desc_lower for word in ['spill', 'discharge', 'pollution']):
                outcomes.append("resulted in environmental impact")
        
        # Build outcomes paragraph
        if outcomes:
            # Use the most specific outcome found
            outcome_text = outcomes[0] if outcomes[0].startswith("resulted") else outcomes[0]
            para = f"The {outcome_text}."
        else:
            para = f"The incident resulted in a marine casualty requiring investigation under 46 CFR Part 4."
        
        # Add response/rescue information if available
        response_entries = [e for e in self.project.timeline if any(word in e.description.lower() 
                           for word in ['coast guard', 'rescue', 'evacuated', 'transported', 'ems', 'medical'])]
        
        if response_entries:
            para += f" {response_entries[0].description}"
        
        return para
    
    def _generate_causal_factors_paragraph(self) -> str:
        """Generate paragraph 3 of Executive Summary - Causal Factors"""
        if not self.project:
            return "Causal factor information not available."
            
        # Identify initiating event
        initiating_event = None
        for entry in self.project.timeline:
            if hasattr(entry, 'is_initiating_event') and entry.is_initiating_event:
                initiating_event = entry
                break
        
        para = "Through its investigation, the Coast Guard determined "
        
        if initiating_event:
            para += f"the initiating event for this casualty was {initiating_event.description.lower()}. "
        else:
            para += f"the cause of this casualty. "
        
        # List causal factors
        if self.project.causal_factors:
            factor_titles = []
            for i, factor in enumerate(self.project.causal_factors):
                factor_titles.append(f"({i+1}) {factor.title}")
            
            para += f"Causal factors contributing to this casualty include: {', '.join(factor_titles)}."
        else:
            para += "Contributing factors are under investigation."
        
        return para
    
    def _describe_incident_from_timeline(self) -> str:
        """Extract incident description from timeline"""
        if not self.project:
            return "an incident occurred"
            
        # Look for event type entries that describe what happened
        for entry in self.project.timeline:
            if entry.type == 'event' or (hasattr(entry, 'is_initiating_event') and entry.is_initiating_event):
                # Clean up the description
                desc = entry.description
                if desc[0].isupper():
                    desc = desc[0].lower() + desc[1:]
                return f"At {self._format_time(entry.timestamp)}, {desc}"
        
        # Fallback
        incident_type = self.project.incident_info.incident_type or "a marine casualty"
        return f"A {incident_type.lower()} occurred"
    
    def _generate_investigating_officers_report(self) -> None:
        """Generate the main Investigating Officer's Report"""
        if not self.document or not self.project:
            return
            
        # [Header block intentionally omitted per new guidance]
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Add date and control number
        p = self.document.add_paragraph()
        p.add_run("16732\n")
        p.add_run(self._format_date(datetime.now()))

        # Add spacing
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Title
        title = self._generate_uscg_title()
        title_para = self.document.add_paragraph()
        title_para.add_run(title).bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()

        # INVESTIGATING OFFICER'S REPORT heading
        report_heading = self.document.add_paragraph()
        report_heading.add_run("INVESTIGATING OFFICER'S REPORT").bold = True
        report_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()

        # Generate all required sections
        self._generate_section_1_preliminary_statement()
        self._generate_section_2_vessels_involved()
        self._generate_section_3_personnel_casualties()
        self._generate_section_4_findings_of_fact()
        self._generate_section_5_analysis()
        self._generate_section_6_conclusions()
        self._generate_section_7_actions_taken()
        self._generate_section_8_recommendations()

        # Add signature block
        self._generate_signature_block()
    
    def _generate_section_1_preliminary_statement(self) -> None:
        """Section 1: Preliminary Statement"""
        if not self.document or not self.project:
            return
            
        # Section heading
        heading = self.document.add_paragraph()
        heading.add_run("1. Preliminary Statement").bold = True

        # 1.1 - Authority statement
        self.document.add_paragraph(
            "1.1. This marine casualty investigation was conducted and this report was submitted in "
            "accordance with Title 46, Code of Federal Regulations (CFR), Subpart 4.07, and under the "
            "authority of Title 46, United States Code (USC) Chapter 63."
        )

        # 1.2 - Parties in interest (if any)
        parties = []
        if self.project.vessels:
            for vessel in self.project.vessels:
                if vessel.official_name:
                    parties.append(f"the owner of the {vessel.official_name}")

        # Add personnel as parties if involved
        for person in self.project.personnel:
            if person.role.lower() in ['master', 'captain', 'operator']:
                parties.append(f"the {person.role.lower()}")

        if parties:
            parties_text = ", ".join(parties)
            self.document.add_paragraph(
                f"1.2. The Investigating Officer designated {parties_text}, as parties-in-interest in this "
                f"investigation. No other individuals, organizations, or parties were designated a party-in-"
                f"interest in accordance with 46 CFR Subsection 4.03-10."
            )

        # 1.3 - Lead agency statement
        self.document.add_paragraph(
            "1.3. The Coast Guard was lead agency for all evidence collection activities involving this "
            "investigation."
        )

        # Check for loss of life
        has_fatality = any('deceased' in e.description.lower() or 'death' in e.description.lower()
                          for e in self.project.timeline)

        if has_fatality:
            self.document.add_paragraph(
                "Due to this investigation involving a loss of life, the Coast Guard Investigative "
                "Service (CGIS) was notified and agreed to provide technical assistance as required."
            )

        self.document.add_paragraph(
            "No other persons or organizations assisted in this investigation."
        )

        # 1.4 - Time format statement
        tz_note = getattr(self.project.incident_info, "time_zone", "local time").replace("_", " ")
        self.document.add_paragraph(
            f"1.4. All times listed in this report are approximate and expressed in {tz_note} using a 24‑hour format."
        )

        self.document.add_paragraph()  # Add spacing
    
    def _generate_section_2_vessels_involved(self) -> None:
        """Section 2: Vessels Involved in the Incident"""
        if not self.document or not self.project:
            return
            
        heading = self.document.add_paragraph()
        heading.add_run("2. Vessels Involved in the Incident").bold = True

        # Get AI-enhanced vessel information
        import logging
        logger = logging.getLogger('app')
        logger.info(f"🟡 ROI SECTION 2: Starting with {len(self.project.vessels)} vessels, {len(self.project.evidence_library)} evidence items")
        enhanced_vessel_info = self._enhance_vessel_information_with_ai()
        logger.info(f"🟡 ROI SECTION 2: Enhanced info keys: {list(enhanced_vessel_info.keys())}")

        for vessel in self.project.vessels:
            # Add photo placeholder
            self.document.add_paragraph()
            photo_para = self.document.add_paragraph()
            photo_para.add_run("Figure 1. Undated Photograph of Vessel").italic = True
            self.document.add_paragraph()

            # Helper for placeholder handling with AI enhancement
            def _safe(val: Any, ai_field: str = None, default: str = "Not documented") -> str:
                import logging
                logger = logging.getLogger('app')
                
                # First try original value
                if val not in [None, "", "##", "YYYY"]:
                    logger.debug(f"🔵 VESSEL SAFE: Using original value for {ai_field}: {val}")
                    return str(val)
                
                # Then try AI-enhanced value
                if ai_field and enhanced_vessel_info.get('vessel_details', {}).get(ai_field):
                    ai_val = enhanced_vessel_info['vessel_details'][ai_field]
                    logger.info(f"🟢 VESSEL SAFE: Using AI value for {ai_field}: {ai_val}")
                    return str(ai_val)
                
                logger.warning(f"⚠️ VESSEL SAFE: Using default for {ai_field}: {default}")
                return default

            # Create vessel information table with comprehensive AI data
            table = self.document.add_table(rows=13, cols=2)
            table.style = 'Table Grid'
            
            # Get AI vessel details
            ai_details = enhanced_vessel_info.get('vessel_details', {})

            # Populate vessel information with all available AI data
            vessel_info = [
                ("Official Name:", _safe(vessel.official_name, 'official_name')),
                ("", f"({_safe(vessel.official_name.upper() if vessel.official_name else None, 'official_name', 'VESSEL NAME').upper()})"),
                ("Identification Number:", self._format_vessel_id(vessel, ai_details)),
                ("Flag:", _safe(ai_details.get('flag') or vessel.flag, default="United States")),
                ("Vessel Class/Type/Sub-Type", self._format_vessel_type(vessel, ai_details)),
                ("Build Year:", _safe(ai_details.get('build_year') or vessel.build_year, 'build_year')),
                ("Gross Tonnage:", self._format_tonnage(vessel, ai_details)),
                ("Length:", self._format_dimensions(vessel, ai_details, 'length')),
                ("Beam/Width:", self._format_dimensions(vessel, ai_details, 'beam')),
                ("Draft/Depth:", self._format_dimensions(vessel, ai_details, 'draft')),
                ("Main/Primary Propulsion:", self._format_propulsion(vessel, ai_details)),
                ("Owner:", self._format_owner_info(vessel, ai_details, 'owner')),
                ("Operator:", self._format_owner_info(vessel, ai_details, 'operator'))
            ]

            # Fill table with vessel information
            for i, (label, value) in enumerate(vessel_info):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Make vessel name italic in second row
                if i == 1:
                    for run in row.cells[1].paragraphs[0].runs:
                        run.italic = True

            self.document.add_paragraph()  # Spacing between vessels

        self.document.add_paragraph()  # Section spacing
    
    def _generate_section_3_personnel_casualties(self) -> None:
        """Section 3: Deceased, Missing, and/or Injured Persons"""
        if not self.document or not self.project:
            return
            
        heading = self.document.add_paragraph()
        heading.add_run("3. Deceased, Missing, and/or Injured Persons").bold = True
        
        # Get AI-enhanced personnel information
        enhanced_personnel_info = self._enhance_personnel_information_with_ai()
        
        # Create table for personnel casualties
        casualties_found = any(p.status.lower() in ['deceased', 'missing', 'injured'] for p in self.project.personnel)
        ai_casualties = enhanced_personnel_info.get('personnel', [])
        ai_casualties_found = any(p.get('status', '').lower() in ['deceased', 'missing', 'injured'] for p in ai_casualties)
        
        if casualties_found or ai_casualties_found:
            table = self.document.add_table(rows=1, cols=4)
            table.style = 'Light Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Relationship to Vessel"
            header_cells[1].text = "Sex"
            header_cells[2].text = "Age"
            header_cells[3].text = "Status"
            
            # Add personnel entries from database
            for person in self.project.personnel:
                if person.status.lower() in ['deceased', 'missing', 'injured']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = person.role or "Unknown"
                    row_cells[1].text = "Unknown"  # Sex not tracked in our model
                    row_cells[2].text = "Unknown"  # Age not tracked in our model
                    row_cells[3].text = person.status.title()
            
            # Add AI-enhanced personnel entries
            for ai_person in ai_casualties:
                if ai_person.get('status', '').lower() in ['deceased', 'missing', 'injured']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ai_person.get('role', 'Unknown')
                    row_cells[1].text = ai_person.get('sex', 'Unknown')
                    row_cells[2].text = str(ai_person.get('age', 'Unknown'))
                    row_cells[3].text = ai_person.get('status', 'Unknown').title()
        else:
            self.document.add_paragraph("No personnel casualties resulted from this incident.")
        
        self.document.add_paragraph()  # Section spacing
    
    def _enhance_vessel_information_with_ai(self) -> Dict[str, Any]:
        """Use AI to extract comprehensive vessel information from evidence files"""
        enhanced_info = {"vessel_details": {}}
        import logging
        logger = logging.getLogger('app')
        
        logger.info("🟡 VESSEL AI: Starting comprehensive vessel information extraction")
        
        try:
            from src.models.anthropic_assistant import AnthropicAssistant
            ai_assistant = AnthropicAssistant()
            
            if not ai_assistant.client:
                logger.error("🔴 VESSEL AI: No Anthropic client available")
                return enhanced_info
            
            # Use the common evidence gathering method
            evidence_content = self._gather_all_evidence_content()
            
            if not evidence_content:
                return enhanced_info
            
            # Comprehensive AI prompt for vessel details
            prompt = f"""
Extract ALL vessel information from this marine casualty investigation evidence. Be extremely thorough.

EVIDENCE CONTENT:
{evidence_content[:25000] if len(evidence_content) > 25000 else evidence_content}

Extract EVERY piece of vessel information mentioned, including:

VESSEL IDENTIFICATION:
- Official name (exactly as written)
- Official number (O.N.)
- Call sign
- IMO number
- State/Federal documentation number
- Port of registry
- Flag state

VESSEL SPECIFICATIONS:
- Type/Class/Sub-type
- Build year
- Builder/Shipyard
- Length (overall, registered, waterline)
- Beam/Width
- Draft/Depth
- Gross tonnage (GT)
- Net tonnage
- Deadweight tonnage
- Passenger capacity
- Crew capacity

PROPULSION & MACHINERY:
- Main engine(s) type, make, model
- Horsepower (total and per engine)
- Propeller configuration
- Auxiliary engines
- Generator specifications
- Steering system
- Fuel capacity and type

OWNERSHIP & OPERATION:
- Owner name and full address
- Operator name and full address
- Managing company
- Charter arrangements
- Insurance company

EQUIPMENT & SYSTEMS:
- Navigation equipment
- Communication equipment
- Safety equipment (life rafts, EPIRBs, etc.)
- Fishing gear (if applicable)
- Cargo handling equipment
- Stability systems

VESSEL HISTORY:
- Previous names
- Previous owners
- Major modifications
- Recent repairs or drydockings
- Inspection history
- Previous incidents

REGULATORY COMPLIANCE:
- Certificate of Inspection status
- Certificate of Documentation
- Safety Management Certificate
- Load line certificate
- Other certificates

Return as comprehensive JSON:
{{
  "vessel_details": {{
    "official_name": "vessel name",
    "official_number": "O.N. number",
    "call_sign": "call sign if found",
    "imo_number": "IMO number if found",
    "documentation_number": "state/federal doc number",
    "port_of_registry": "home port",
    "flag": "flag state",
    "vessel_class": "vessel classification",
    "vessel_type": "specific type",
    "vessel_subtype": "sub-type if applicable",
    "build_year": "year built",
    "builder": "shipyard/builder name",
    "length": "length in feet",
    "length_type": "overall/registered/waterline",
    "beam": "beam in feet",
    "draft": "draft in feet",
    "gross_tonnage": "GT value",
    "net_tonnage": "NT value",
    "deadweight_tonnage": "DWT if applicable",
    "passenger_capacity": "number if applicable",
    "crew_capacity": "number",
    "propulsion": "detailed engine description",
    "main_engine_make": "manufacturer",
    "main_engine_model": "model",
    "horsepower": "total HP",
    "propellers": "number and type",
    "auxiliary_engines": "description",
    "generators": "description",
    "fuel_capacity": "gallons/liters",
    "fuel_type": "diesel/gasoline/other",
    "owner": "full owner name",
    "owner_address": "complete address",
    "operator": "full operator name",
    "operator_address": "complete address",
    "managing_company": "if different from owner/operator",
    "navigation_equipment": "list of nav equipment",
    "communication_equipment": "list of comm equipment",
    "safety_equipment": "list of safety gear",
    "certificates": "list of certificates and expiration dates",
    "previous_names": "if mentioned",
    "modifications": "major modifications",
    "last_drydock": "date and location",
    "inspection_history": "recent inspections"
  }}
}}

Extract EVERYTHING mentioned. Use null for fields not found in the evidence.
"""
            
            logger.info("🟡 VESSEL AI: Sending comprehensive prompt to AI assistant")
            response = ai_assistant.chat(prompt)
            
            # Use safe JSON extraction
            raw_enhanced_info = ai_assistant._safe_json_extract(response)
            
            if isinstance(raw_enhanced_info, dict) and 'vessel_details' in raw_enhanced_info:
                enhanced_info = raw_enhanced_info
                logger.info(f"🟢 VESSEL AI: Successfully extracted comprehensive vessel information")
            else:
                logger.warning(f"⚠️ VESSEL AI: Unexpected response structure")
            
        except Exception as e:
            logger.error(f"🔴 VESSEL AI: Error extracting vessel info: {e}")
            enhanced_info = {"vessel_details": {}}
        
        return enhanced_info
    
    def _enhance_personnel_information_with_ai(self) -> Dict[str, Any]:
        """Use AI to extract comprehensive personnel information from evidence files"""
        enhanced_info = {"personnel": []}
        import logging
        logger = logging.getLogger('app')
        
        logger.info("🟡 PERSONNEL AI: Starting comprehensive personnel extraction")
        
        try:
            from src.models.anthropic_assistant import AnthropicAssistant
            ai_assistant = AnthropicAssistant()
            
            if not ai_assistant.client:
                return enhanced_info
            
            # Use the common evidence gathering method
            evidence_content = self._gather_all_evidence_content()
            
            if not evidence_content:
                return enhanced_info
            
            # Comprehensive AI prompt for personnel details
            prompt = f"""
Extract ALL personnel information from this marine casualty investigation evidence. Be extremely thorough.

EVIDENCE CONTENT:
{evidence_content[:25000] if len(evidence_content) > 25000 else evidence_content}

Extract information about EVERY person mentioned, including:

IDENTIFICATION:
- Full name (exactly as written)
- Alternative names/nicknames
- Date of birth
- Age at time of incident
- Sex/Gender
- Nationality
- Home address

PROFESSIONAL DETAILS:
- Current role/position
- Vessel assignment
- Employment status (permanent/temporary/contract)
- Years of experience (total and in current role)
- Previous positions
- Employer/Company

CREDENTIALS & QUALIFICATIONS:
- License type and number
- License issuing authority
- License expiration date
- Endorsements
- Training certificates
- Medical certificates
- Drug test results

INCIDENT INVOLVEMENT:
- Location during incident
- Actions taken
- Injuries sustained (detailed)
- Medical treatment received
- Hospital/medical facility
- Time of death (if applicable)
- Cause of death (if applicable)

BACKGROUND:
- Previous incidents/violations
- Performance history
- Recent work schedule
- Rest periods before incident
- Physical/mental condition

For EACH person mentioned, create an entry:
{{
  "personnel": [
    {{
      "name": "Full name as written",
      "alternative_names": "Other names/nicknames",
      "date_of_birth": "DOB if mentioned",
      "age": "Age at incident",
      "sex": "Male/Female/Unknown",
      "nationality": "Country",
      "home_address": "Full address if available",
      "role": "Specific position/title",
      "vessel_assignment": "Vessel name",
      "employment_status": "Permanent/Temporary/Contract",
      "employer": "Company name",
      "years_experience_total": "Total maritime experience",
      "years_experience_role": "Experience in current position",
      "license_type": "Master/Mate/Engineer/etc",
      "license_number": "License number",
      "license_authority": "USCG/Other",
      "license_expiration": "Expiration date",
      "endorsements": "List of endorsements",
      "certificates": "Training/medical certificates",
      "status": "Deceased/Injured/Missing/Uninjured",
      "injuries": "Detailed injury description",
      "medical_treatment": "Treatment received",
      "hospital": "Medical facility name",
      "time_of_death": "If applicable",
      "cause_of_death": "If applicable",
      "location_during_incident": "Where they were",
      "actions_during_incident": "What they did",
      "drug_test_results": "Positive/Negative/Pending",
      "alcohol_test_results": "BAC level if tested",
      "previous_incidents": "Any prior violations/incidents",
      "work_schedule": "Recent duty hours",
      "rest_period": "Hours of rest before incident",
      "additional_details": "Any other relevant information"
    }}
  ]
}}

Include EVERYONE mentioned: crew, passengers, responders, medical personnel, investigators, witnesses.
Extract ALL available information. Use null for fields not found in evidence.
"""
            
            logger.info("🟡 PERSONNEL AI: Sending comprehensive prompt to AI assistant")
            response = ai_assistant.chat(prompt)
            
            # Use safe JSON extraction
            raw_enhanced_info = ai_assistant._safe_json_extract(response)
            
            if isinstance(raw_enhanced_info, dict) and 'personnel' in raw_enhanced_info:
                enhanced_info = raw_enhanced_info
                logger.info(f"🟢 PERSONNEL AI: Extracted {len(enhanced_info['personnel'])} personnel records")
            else:
                logger.warning(f"⚠️ PERSONNEL AI: Unexpected response structure")
            
        except Exception as e:
            logger.error(f"🔴 PERSONNEL AI: Error extracting personnel info: {e}")
            enhanced_info = {"personnel": []}
        
        return enhanced_info

    def _generate_section_4_findings_of_fact(self) -> None:
        """Section 4: Findings of Fact"""
        if not self.document or not self.project:
            return
            
        heading = self.document.add_paragraph()
        heading.add_run("4. Findings of Fact").bold = True
        
        # 4.1 The Incident
        subheading = self.document.add_paragraph()
        subheading.add_run("4.1. The Incident:").bold = True
        
        # Generate professional findings from timeline using AI - ENHANCED VERSION
        if hasattr(self.project, 'roi_document') and self.project.roi_document.findings_of_fact:
            # Use existing findings if available (statements are stored without numbers)
            finding_number = 1
            for finding in self.project.roi_document.findings_of_fact:
                para = self.document.add_paragraph(f"4.1.{finding_number}. {finding.statement}")
                finding_number += 1
        else:
            # Generate comprehensive findings from timeline using Anthropic AI
            from src.models.anthropic_assistant import AnthropicAssistant
            anthropic_assistant = AnthropicAssistant()
            
            if anthropic_assistant.client and self.project.timeline:
                # Convert timeline entries to appropriate format
                timeline_objects: List[TimelineEntry] = []
                for entry in self.project.timeline:
                    timeline_objects.append(entry)
                
                evidence_objects: List[Evidence] = []
                for evidence in self.project.evidence_library:
                    evidence_objects.append(evidence)
                
                # Generate professional findings using Anthropic
                findings_statements = anthropic_assistant.generate_findings_of_fact_from_timeline(timeline_objects, evidence_objects)
                
                # Add AI-generated findings (they already have proper numbering)
                for finding_statement in findings_statements:
                    # AI assistants already provide properly numbered findings
                    finding_text = finding_statement.strip()
                    para = self.document.add_paragraph(finding_text)
                
                # If no AI findings generated, use enhanced fallback
                if not findings_statements:
                    self._generate_enhanced_findings_from_timeline()
            else:
                # Enhanced fallback method
                self._generate_enhanced_findings_from_timeline()
        
        # 4.2 Additional/Supporting Information (if needed)
        self._generate_section_4_2_supporting_findings()
        
        self.document.add_paragraph()  # Section spacing
    
    def _generate_section_4_2_supporting_findings(self) -> None:
        """Generate Section 4.2 - Supporting/Background Information"""
        if not self.document or not self.project:
            return
            
        # Identify incident date from initiating event
        incident_date: Optional[date] = None
        for entry in self.project.timeline:
            if hasattr(entry, 'is_initiating_event') and entry.is_initiating_event and entry.timestamp:
                incident_date = entry.timestamp.date()
                break
        
        # Collect background timeline entries (pre-incident and post-casualty)
        background_entries: List[TimelineEntry] = []
        post_casualty_entries: List[TimelineEntry] = []
        
        if incident_date:
            for entry in self.project.timeline:
                if entry.timestamp:
                    if entry.timestamp.date() < incident_date:
                        background_entries.append(entry)
                    elif entry.timestamp.date() > incident_date:
                        post_casualty_entries.append(entry)
        
        # Sort entries
        background_entries.sort(key=lambda x: x.timestamp or datetime.min)
        post_casualty_entries.sort(key=lambda x: x.timestamp or datetime.min)
        
        # Only add section if there's supporting information
        if not (background_entries or post_casualty_entries or self.project.evidence_library):
            return
        
        self.document.add_paragraph()
        subheading = self.document.add_paragraph()
        subheading.add_run("4.2. Additional/Supporting Information:").bold = True
        
        finding_number = 1
        
        # Pre-incident background findings
        if background_entries:
            # Group by general timeframe to avoid too many findings
            for entry in background_entries[-5:]:  # Last 5 most relevant background entries
                time_str = self._format_date(entry.timestamp)
                description = entry.description
                if not description.endswith('.'):
                    description += '.'
                
                # Format as background finding
                finding_text = f"4.2.{finding_number}. Prior to the incident, on {time_str}, {description}"
                self.document.add_paragraph(finding_text)
                finding_number += 1
        
        # Add evidence-based findings if AI not available
        if self.project.evidence_library and finding_number <= 3:
            # Try to extract meaningful information from evidence
            from src.models.anthropic_assistant import AnthropicAssistant
            anthropic_assistant = AnthropicAssistant()
            
            if anthropic_assistant.client:
                try:
                    # Generate background findings from evidence - fix the date parameter
                    evidence_findings = anthropic_assistant.generate_background_findings_from_evidence(
                        self.project.evidence_library,
                        incident_date
                    )
                    
                    for finding in evidence_findings[:5]:  # Limit to 5 findings
                        finding_text = f"4.2.{finding_number}. {finding}"
                        self.document.add_paragraph(finding_text)
                        finding_number += 1
                except Exception as e:
                    print(f"Error generating background findings: {e}")
                    pass
            
            # Fallback if AI not available
            if finding_number == 1:
                for i, evidence in enumerate(self.project.evidence_library[:3], finding_number):
                    if evidence.type in ['crew_statement', 'inspection_report', 'maintenance_record']:
                        finding_text = f"4.2.{i}. Documentary evidence '{evidence.filename}' provides {evidence.type.replace('_', ' ')} regarding vessel operations."
                        self.document.add_paragraph(finding_text)
                        finding_number += 1
        
        # Post-casualty findings
        if post_casualty_entries and finding_number <= 8:
            for entry in post_casualty_entries[:3]:  # First 3 post-casualty entries
                time_str = self._format_date(entry.timestamp) + f", at {self._format_time(entry.timestamp)}"
                description = entry.description
                if not description.endswith('.'):
                    description += '.'
                
                finding_text = f"4.2.{finding_number}. Following the casualty, on {time_str}, {description}"
                self.document.add_paragraph(finding_text)
                finding_number += 1
    
    def _generate_enhanced_findings_from_timeline(self) -> None:
        """Enhanced method for professional timeline to findings conversion - focuses on incident day"""
        if not self.document or not self.project:
            return
            
        if not self.project.timeline:
            self.document.add_paragraph("4.1.1. No timeline entries have been documented for this incident.")
            return
        
        # Identify incident date from initiating event
        incident_date: Optional[date] = None
        for entry in self.project.timeline:
            if hasattr(entry, 'is_initiating_event') and entry.is_initiating_event and entry.timestamp:
                incident_date = entry.timestamp.date()
                break
        
        # Filter for incident-day entries only
        incident_entries: List[TimelineEntry] = []
        if incident_date:
            incident_entries = [
                entry for entry in self.project.timeline 
                if entry.timestamp and entry.timestamp.date() == incident_date
            ]
        else:
            # Fallback: use all entries if no incident date identified
            incident_entries = [entry for entry in self.project.timeline if entry.timestamp]
        
        # Sort timeline by timestamp
        sorted_timeline = sorted(incident_entries, key=lambda x: x.timestamp or datetime.min)
        
        if not sorted_timeline:
            self.document.add_paragraph("4.1.1. No incident-day timeline entries have been documented.")
            return
        
        finding_number = 1
        for entry in sorted_timeline:
            # Format timestamp in USCG style
            time_str = self._format_date(entry.timestamp) + f", at {self._format_time(entry.timestamp)}"
            
            # Create professional finding statement with enhanced formatting
            description = entry.description
            if not description.endswith('.'):
                description += '.'
            
            # Capitalize first letter for professional presentation
            if description and description[0].islower():
                description = description[0].upper() + description[1:]
            
            finding_text = f"4.1.{finding_number}. On {time_str}, {description}"
            self.document.add_paragraph(finding_text)
            finding_number += 1
    
    def _generate_section_5_analysis(self) -> None:
        """Section 5: Analysis - THE MOST CRITICAL SECTION OF THE ROI"""
        if not self.document or not self.project:
            return
            
        # Add emphasis that this is the most important section
        heading = self.document.add_paragraph()
        heading.add_run("5. Analysis").bold = True
        
        # Brief professional introduction
        intro_para = self.document.add_paragraph()
        intro_para.add_run("The Coast Guard's investigation identified the following causal factors that contributed to this marine casualty:")
        self.document.add_paragraph()  # Extra spacing
        
        # Check if we have causal factors to analyze
        if not self.project.causal_factors:
            # Generate a placeholder analysis section if no factors exist
            self.document.add_paragraph("No causal factors have been identified for analysis at this time. Further investigation may be required to determine contributing factors.")
            self.document.add_paragraph()
            return
        
        # Generate analysis for each causal factor with enhanced formatting
        analysis_number = 1
        for factor in self.project.causal_factors:
            # Analysis heading with negative phrasing emphasis
            subheading = self.document.add_paragraph()
            
            # Ensure negative phrasing in title
            title = factor.title
            if not any(title.lower().startswith(phrase) for phrase in ['failure', 'lack', 'inadequate', 'absence', 'deficiency']):
                if 'failure' not in title.lower() and 'lack' not in title.lower():
                    title = f"Failure to properly address: {title}"
            
            subheading_run = subheading.add_run(f"5.{analysis_number}. {title}")
            subheading_run.bold = True
            subheading_run.font.size = Pt(12)
            
            # Use Anthropic to generate concise professional analysis
            from src.models.anthropic_assistant import AnthropicAssistant
            anthropic_assistant = AnthropicAssistant()
            
            analysis_para = self.document.add_paragraph()
            
            # Get improved analysis text from Anthropic
            if anthropic_assistant.client:
                try:
                    analysis_text = anthropic_assistant.improve_analysis_text(factor)
                except Exception as e:
                    print(f"Error improving analysis text: {e}")
                    # Fallback to simple format if Anthropic fails
                    analysis_text = factor.analysis_text or factor.description
                    if not analysis_text.lower().startswith('it is reasonable'):
                        if factor.category == 'precondition':
                            analysis_text = f"It is reasonable to believe that {analysis_text.lower()} contributed to the casualty sequence."
                        elif factor.category == 'production':
                            analysis_text = f"The {analysis_text.lower()} was a direct factor in the incident."
                        elif factor.category == 'defense':
                            analysis_text = f"The absence of {analysis_text.lower()} allowed the casualty to occur."
                        else:
                            analysis_text = f"It is reasonable to believe that {analysis_text.lower()} was a contributing factor."
            else:
                # Fallback to simple format if Anthropic not available
                analysis_text = factor.analysis_text or factor.description
                if not analysis_text.lower().startswith('it is reasonable'):
                    if factor.category == 'precondition':
                        analysis_text = f"It is reasonable to believe that {analysis_text.lower()} contributed to the casualty sequence."
                    elif factor.category == 'production':
                        analysis_text = f"The {analysis_text.lower()} was a direct factor in the incident."
                    elif factor.category == 'defense':
                        analysis_text = f"The absence of {analysis_text.lower()} allowed the casualty to occur."
                    else:
                        analysis_text = f"It is reasonable to believe that {analysis_text.lower()} was a contributing factor."
            
            analysis_para.add_run(analysis_text)
            
            # Add spacing between analyses for clarity
            self.document.add_paragraph()
            analysis_number += 1
    
        self.document.add_paragraph()  # Section spacing
    
    def _generate_section_6_conclusions(self) -> None:
        """Section 6: Conclusions – using AI to extract from evidence"""
        if not self.document or not self.project:
            return
            
        heading = self.document.add_paragraph()
        heading.add_run("6. Conclusions").bold = True
        
        # Use AI to generate comprehensive conclusions from evidence
        import logging
        logger = logging.getLogger('app')
        logger.info("🟡 CONCLUSIONS: Generating AI-based conclusions from evidence")
        
        from src.models.anthropic_assistant import AnthropicAssistant
        ai_assistant = AnthropicAssistant()
        
        if ai_assistant.client:
            conclusions_data = self._generate_conclusions_with_ai()
            
            if conclusions_data:
                # 6.1 Determination of Cause
                subheading = self.document.add_paragraph()
                subheading.add_run("6.1. Determination of Cause:").bold = True
                
                # Add initiating event conclusion
                if conclusions_data.get('initiating_event'):
                    self.document.add_paragraph(f"6.1.1. {conclusions_data['initiating_event']}")
                else:
                    self.document.add_paragraph(
                        "6.1.1. The initiating event for this casualty could not be conclusively determined; "
                        "however, the following scenario is considered most probable based on available evidence."
                    )
                
                # Add causal factor conclusions
                if conclusions_data.get('causal_determinations'):
                    for i, determination in enumerate(conclusions_data['causal_determinations'], 1):
                        self.document.add_paragraph(f"6.1.1.{i}. {determination}")
                
                # Add other conclusion sections from AI analysis
                for section_num in ['6.2', '6.3', '6.4', '6.5', '6.6']:
                    if conclusions_data.get(f'section_{section_num}'):
                        self.document.add_paragraph(conclusions_data[f'section_{section_num}'])
                    else:
                        # Use default if AI didn't provide specific content
                        if section_num == '6.2':
                            self.document.add_paragraph("6.2. Evidence of Act(s) or Violation(s) of Law by Credentialed Mariners: None identified.")
                        elif section_num == '6.3':
                            self.document.add_paragraph("6.3. Evidence of Act(s) or Violation(s) of Law by U.S. Coast Guard personnel or others: None identified.")
                        elif section_num == '6.4':
                            self.document.add_paragraph("6.4. Evidence of Act(s) Subject to Civil Penalty: None identified.")
                        elif section_num == '6.5':
                            self.document.add_paragraph("6.5. Evidence of Criminal Act(s): None identified.")
                        elif section_num == '6.6':
                            self.document.add_paragraph("6.6. Need for New or Amended U.S. Law or Regulation: None identified.")
            else:
                # Fallback to original logic if AI fails
                self._generate_conclusions_fallback()
        else:
            # Fallback to original logic if no AI
            self._generate_conclusions_fallback()

        self.document.add_paragraph()  # Section spacing
    
    def _generate_conclusions_fallback(self) -> None:
        """Fallback method for conclusions if AI is unavailable"""
        # Original logic preserved as fallback
        subheading = self.document.add_paragraph()
        subheading.add_run("6.1. Determination of Cause:").bold = True

        initiating_event = next(
            (e for e in self.project.timeline if getattr(e, "is_initiating_event", False)), None
        )

        if initiating_event and initiating_event.description:
            init_text = (
                f"6.1.1. The initiating event for this casualty occurred when "
                f"{initiating_event.description.lower()}."
            )
        else:
            init_text = (
                "6.1.1. The initiating event for this casualty could not be conclusively determined; "
                "however, the following scenario is considered most probable based on available evidence."
            )
        self.document.add_paragraph(init_text)

        if self.project.causal_factors:
            sub_index = 1
            for factor in self.project.causal_factors:
                text = factor.description or factor.title
                para_num = f"6.1.1.{sub_index}"
                self.document.add_paragraph(f"{para_num}. {text.rstrip('.')}." )
                sub_index += 1
        else:
            self.document.add_paragraph(
                "6.1.1.1. No specific causal factors were identified during the investigation."
            )

        skip_lines = [
            "6.2. Evidence of Act(s) or Violation(s) of Law by Credentialed Mariners: None identified.",
            "6.3. Evidence of Act(s) or Violation(s) of Law by U.S. Coast Guard personnel or others: None identified.",
            "6.4. Evidence of Act(s) Subject to Civil Penalty: None identified.",
            "6.5. Evidence of Criminal Act(s): None identified.",
            "6.6. Need for New or Amended U.S. Law or Regulation: None identified.",
        ]
        for line in skip_lines:
            self.document.add_paragraph(line)
    
    def _generate_section_7_actions_taken(self) -> None:
        """Section 7: Actions Taken Since the Incident - using AI to extract from evidence"""
        if not self.document or not self.project:
            return
            
        heading = self.document.add_paragraph()
        heading.add_run("7. Actions Taken Since the Incident").bold = True
        
        # Use AI to extract actions taken from evidence documents
        import logging
        logger = logging.getLogger('app')
        logger.info("🟡 ACTIONS: Generating AI-based actions taken from evidence")
        
        from src.models.anthropic_assistant import AnthropicAssistant
        ai_assistant = AnthropicAssistant()
        
        if ai_assistant.client:
            actions_taken = self._generate_actions_taken_with_ai()
            
            if actions_taken and len(actions_taken) > 0:
                for i, action in enumerate(actions_taken, 1):
                    self.document.add_paragraph(f"7.{i}. {action}")
            else:
                # If AI doesn't find specific actions, add minimal default
                self.document.add_paragraph(
                    "7.1. The Coast Guard initiated a formal investigation under 46 CFR Part 4."
                )
        else:
            # Fallback if no AI available
            self.document.add_paragraph(
                "7.1. The Coast Guard initiated a formal investigation under 46 CFR Part 4."
            )
        
        self.document.add_paragraph()  # Section spacing
    
    def _generate_section_8_recommendations(self) -> None:
        """Section 8: Recommendations - using AI to generate from evidence and analysis"""
        if not self.document or not self.project:
            return
            
        heading = self.document.add_paragraph()
        heading.add_run("8. Recommendations").bold = True
        
        # Use AI to generate comprehensive recommendations
        import logging
        logger = logging.getLogger('app')
        logger.info("🟡 RECOMMENDATIONS: Generating AI-based recommendations from evidence and analysis")
        
        from src.models.anthropic_assistant import AnthropicAssistant
        ai_assistant = AnthropicAssistant()
        
        if ai_assistant.client:
            recommendations_data = self._generate_recommendations_with_ai()
            
            if recommendations_data:
                # 8.1 Safety Recommendations
                if recommendations_data.get('safety_recommendations'):
                    subheading = self.document.add_paragraph()
                    subheading.add_run("8.1. Safety Recommendations:").bold = True
                    
                    for i, rec in enumerate(recommendations_data['safety_recommendations'], 1):
                        self.document.add_paragraph(f"8.1.{i}. {rec}")
                
                # 8.2 Administrative Recommendations
                self.document.add_paragraph()
                subheading = self.document.add_paragraph()
                subheading.add_run("8.2. Administrative Recommendations:").bold = True
                
                if recommendations_data.get('administrative_recommendations'):
                    for i, rec in enumerate(recommendations_data['administrative_recommendations'], 1):
                        self.document.add_paragraph(f"8.2.{i}. {rec}")
                else:
                    self.document.add_paragraph("None at this time.")
            else:
                # Minimal fallback if AI fails
                self._generate_recommendations_fallback()
        else:
            # Fallback if no AI available
            self._generate_recommendations_fallback()
        
        self.document.add_paragraph()  # Section spacing
    
    def _generate_recommendations_fallback(self) -> None:
        """Fallback recommendations if AI is unavailable"""
        subheading = self.document.add_paragraph()
        subheading.add_run("8.1. Safety Recommendations:").bold = True
        self.document.add_paragraph(
            "8.1.1. Conduct a comprehensive review of vessel safety procedures and emergency response protocols."
        )
        
        self.document.add_paragraph()
        subheading = self.document.add_paragraph()
        subheading.add_run("8.2. Administrative Recommendations:").bold = True
        self.document.add_paragraph("None at this time.")
    
    def _generate_signature_block(self) -> None:
        """Generate signature block"""
        if not self.document or not self.project:
            return
            
        # Add spacing
        for _ in range(3):
            self.document.add_paragraph()
        
        # Signature line
        self.document.add_paragraph("_" * 50)
        
        # Name and title
        officer_name = self.project.metadata.investigating_officer or "[Investigating Officer Name]"
        self.document.add_paragraph(officer_name.upper())
        self.document.add_paragraph("Lieutenant, U.S. Coast Guard")
        self.document.add_paragraph("Investigating Officer")
    
    def _generate_conclusions_with_ai(self) -> Dict[str, Any]:
        """Use AI to extract conclusions from all evidence"""
        import logging
        logger = logging.getLogger('app')
        
        try:
            from src.models.anthropic_assistant import AnthropicAssistant
            ai_assistant = AnthropicAssistant()
            
            if not ai_assistant.client:
                return {}
            
            # Gather all evidence content
            evidence_content = self._gather_all_evidence_content()
            
            prompt = f"""
Analyze this marine casualty investigation evidence to generate USCG Section 6 Conclusions.

EVIDENCE CONTENT:
{evidence_content[:20000] if len(evidence_content) > 20000 else evidence_content}

Generate professional conclusions following USCG format:

1. INITIATING EVENT (6.1.1): Identify the first adverse outcome that started the casualty sequence
2. CAUSAL DETERMINATIONS (6.1.1.1, 6.1.1.2, etc.): List specific causal factors that contributed
3. VIOLATIONS BY MARINERS (6.2): Any evidence of violations by credentialed mariners
4. VIOLATIONS BY USCG/OTHERS (6.3): Any evidence of violations by Coast Guard or other personnel
5. CIVIL PENALTY EVIDENCE (6.4): Any acts subject to civil penalties
6. CRIMINAL ACTS (6.5): Any evidence of criminal activity
7. REGULATORY NEEDS (6.6): Any need for new or amended regulations

Return as JSON:
{{
  "initiating_event": "The initiating event for this casualty was...",
  "causal_determinations": [
    "Factor 1 description",
    "Factor 2 description"
  ],
  "section_6.2": "6.2. Evidence of Act(s) or Violation(s)...: [Specific findings or 'None identified']",
  "section_6.3": "6.3. Evidence of Act(s) or Violation(s)...: [Specific findings or 'None identified']",
  "section_6.4": "6.4. Evidence of Act(s) Subject to Civil Penalty: [Specific findings or 'None identified']",
  "section_6.5": "6.5. Evidence of Criminal Act(s): [Specific findings or 'None identified']",
  "section_6.6": "6.6. Need for New or Amended U.S. Law or Regulation: [Specific findings or 'None identified']"
}}
"""
            
            response = ai_assistant.chat(prompt)
            conclusions_data = ai_assistant._safe_json_extract(response)
            
            logger.info(f"🟢 CONCLUSIONS AI: Successfully extracted conclusions from evidence")
            return conclusions_data
            
        except Exception as e:
            logger.error(f"🔴 CONCLUSIONS AI: Error generating conclusions: {e}")
            return {}
    
    def _generate_actions_taken_with_ai(self) -> List[str]:
        """Use AI to extract actions taken from evidence"""
        import logging
        logger = logging.getLogger('app')
        
        try:
            from src.models.anthropic_assistant import AnthropicAssistant
            ai_assistant = AnthropicAssistant()
            
            if not ai_assistant.client:
                return []
            
            # Gather all evidence content
            evidence_content = self._gather_all_evidence_content()
            
            prompt = f"""
Analyze this marine casualty investigation evidence to identify all actions taken since the incident.

EVIDENCE CONTENT:
{evidence_content[:20000] if len(evidence_content) > 20000 else evidence_content}

Extract ALL actions taken by:
- Coast Guard (investigations, testing, orders, notifications)
- Vessel operators/owners (repairs, policy changes, training)
- Other agencies (medical response, environmental cleanup)
- Industry organizations (safety bulletins, guidance)

Focus on POST-INCIDENT actions only. Include:
- Drug/alcohol testing conducted
- Captain of the Port orders issued
- Safety notifications distributed
- Vessel inspections performed
- Policy or procedure changes
- Training conducted
- Equipment repairs or replacements
- Regulatory enforcement actions

Return as JSON array of action statements:
[
  "The Coast Guard conducted post-casualty drug and alcohol testing...",
  "A Captain of the Port order was issued requiring...",
  "The vessel operator implemented new safety procedures..."
]

Each action should be a complete, professional statement.
"""
            
            response = ai_assistant.chat(prompt)
            actions = ai_assistant._safe_json_extract(response)
            
            if isinstance(actions, list):
                logger.info(f"🟢 ACTIONS AI: Extracted {len(actions)} actions from evidence")
                return actions
            else:
                logger.warning("⚠️ ACTIONS AI: Response was not a list")
                return []
                
        except Exception as e:
            logger.error(f"🔴 ACTIONS AI: Error extracting actions: {e}")
            return []
    
    def _generate_recommendations_with_ai(self) -> Dict[str, Any]:
        """Use AI to generate recommendations based on entire investigation"""
        import logging
        logger = logging.getLogger('app')
        
        try:
            from src.models.anthropic_assistant import AnthropicAssistant
            ai_assistant = AnthropicAssistant()
            
            if not ai_assistant.client:
                return {}
            
            # Gather all evidence content and causal factors
            evidence_content = self._gather_all_evidence_content()
            causal_summary = self._summarize_causal_factors()
            
            prompt = f"""
Based on this marine casualty investigation, generate comprehensive recommendations to prevent similar incidents.

INCIDENT SUMMARY:
{self.project.incident_info.incident_type} at {self.project.incident_info.location}

CAUSAL FACTORS IDENTIFIED:
{causal_summary}

EVIDENCE REVIEWED:
{evidence_content[:15000] if len(evidence_content) > 15000 else evidence_content}

Generate recommendations in two categories:

SAFETY RECOMMENDATIONS (8.1):
- Vessel-specific improvements (equipment, procedures, training)
- Industry-wide safety enhancements
- Regulatory compliance improvements
- Emergency response enhancements
- Communication and coordination improvements

ADMINISTRATIVE RECOMMENDATIONS (8.2):
- Policy changes
- Documentation requirements
- Inspection or audit programs
- Enforcement actions
- Inter-agency coordination

Each recommendation should:
- Address specific causal factors identified
- Be actionable and specific
- Follow USCG professional format
- Include who should implement it

Return as JSON:
{{
  "safety_recommendations": [
    "Vessel operators should implement...",
    "The maritime industry should develop...",
    "Training programs should include..."
  ],
  "administrative_recommendations": [
    "The Coast Guard should consider...",
    "Marine inspectors should verify...",
    "Policy guidance should be updated..."
  ]
}}

If no administrative recommendations are warranted, return empty array.
"""
            
            response = ai_assistant.chat(prompt)
            recommendations = ai_assistant._safe_json_extract(response)
            
            logger.info(f"🟢 RECOMMENDATIONS AI: Generated recommendations from investigation")
            return recommendations
            
        except Exception as e:
            logger.error(f"🔴 RECOMMENDATIONS AI: Error generating recommendations: {e}")
            return {}
    
    def _gather_all_evidence_content(self) -> str:
        """Gather all evidence content from uploaded files"""
        import logging
        logger = logging.getLogger('app')
        
        evidence_content = ""
        evidence_count = 0
        
        logger.info(f"🟡 EVIDENCE GATHER: Processing {len(self.project.evidence_library)} evidence items")
        
        for evidence in self.project.evidence_library:
            try:
                if hasattr(evidence, 'file_path') and evidence.file_path:
                    import os
                    from flask import current_app
                    uploads_dir = current_app.config.get('UPLOAD_FOLDER', 'uploads')
                    file_path = os.path.join(uploads_dir, evidence.file_path)
                    
                    if os.path.exists(file_path):
                        from src.models.project_manager import ProjectManager
                        pm = ProjectManager()
                        content = pm._extract_file_content(file_path)
                        if content:
                            evidence_content += f"\n\n--- DOCUMENT: {evidence.filename} ---\n"
                            evidence_content += content
                            evidence_count += 1
                elif hasattr(evidence, 'content') and evidence.content:
                    evidence_content += f"\n\n--- EVIDENCE: {evidence.type} ---\n"
                    evidence_content += evidence.content
                    evidence_count += 1
            except Exception as e:
                logger.warning(f"⚠️ EVIDENCE GATHER: Error processing {evidence.filename}: {e}")
                continue
        
        logger.info(f"🟢 EVIDENCE GATHER: Collected content from {evidence_count} evidence items")
        return evidence_content
    
    def _summarize_causal_factors(self) -> str:
        """Summarize causal factors for recommendation generation"""
        if not self.project.causal_factors:
            return "No specific causal factors were identified."
        
        summary = []
        for factor in self.project.causal_factors:
            summary.append(f"- {factor.category.upper()}: {factor.title}")
            if factor.description:
                summary.append(f"  Description: {factor.description}")
        
        return "\n".join(summary)
    
    def _format_vessel_id(self, vessel, ai_details: Dict[str, Any]) -> str:
        """Format vessel identification with all available numbers"""
        parts = []
        
        # Official number
        if ai_details.get('official_number'):
            parts.append(f"{ai_details['official_number']} - Official Number (US)")
        elif vessel.identification_number:
            parts.append(f"{vessel.identification_number} - Official Number (US)")
        
        # Call sign
        if ai_details.get('call_sign'):
            parts.append(f"Call Sign: {ai_details['call_sign']}")
        
        # IMO number
        if ai_details.get('imo_number'):
            parts.append(f"IMO: {ai_details['imo_number']}")
        
        # Documentation number
        if ai_details.get('documentation_number'):
            parts.append(f"Doc: {ai_details['documentation_number']}")
        
        return "; ".join(parts) if parts else "Not documented"
    
    def _format_vessel_type(self, vessel, ai_details: Dict[str, Any]) -> str:
        """Format vessel type information"""
        vessel_class = ai_details.get('vessel_class') or vessel.vessel_class or "Unknown"
        vessel_type = ai_details.get('vessel_type') or vessel.vessel_type or "Unknown"
        vessel_subtype = ai_details.get('vessel_subtype') or vessel.vessel_subtype or "N/A"
        
        return f"{vessel_class}/{vessel_type}/{vessel_subtype}"
    
    def _format_tonnage(self, vessel, ai_details: Dict[str, Any]) -> str:
        """Format tonnage information"""
        parts = []
        
        # Gross tonnage
        gt = ai_details.get('gross_tonnage') or vessel.gross_tonnage
        if gt:
            parts.append(f"{gt} GT")
        
        # Net tonnage
        if ai_details.get('net_tonnage'):
            parts.append(f"{ai_details['net_tonnage']} NT")
        
        # Deadweight
        if ai_details.get('deadweight_tonnage'):
            parts.append(f"{ai_details['deadweight_tonnage']} DWT")
        
        return "; ".join(parts) if parts else "Not documented"
    
    def _format_dimensions(self, vessel, ai_details: Dict[str, Any], dimension: str) -> str:
        """Format vessel dimensions with type information"""
        value = ai_details.get(dimension) or getattr(vessel, dimension, None)
        
        if not value:
            return "Not documented"
        
        # Add dimension type if available
        if dimension == 'length' and ai_details.get('length_type'):
            return f"{value} feet ({ai_details['length_type']})"
        else:
            return f"{value} feet"
    
    def _format_propulsion(self, vessel, ai_details: Dict[str, Any]) -> str:
        """Format comprehensive propulsion information"""
        parts = []
        
        # Main engine details
        if ai_details.get('main_engine_make') and ai_details.get('main_engine_model'):
            engine = f"{ai_details['main_engine_make']} {ai_details['main_engine_model']}"
            if ai_details.get('horsepower'):
                engine += f", {ai_details['horsepower']} HP"
            parts.append(engine)
        elif ai_details.get('propulsion'):
            parts.append(ai_details['propulsion'])
        elif vessel.propulsion:
            parts.append(vessel.propulsion)
        
        # Propellers
        if ai_details.get('propellers'):
            parts.append(f"Propellers: {ai_details['propellers']}")
        
        # Fuel type
        if ai_details.get('fuel_type'):
            parts.append(f"Fuel: {ai_details['fuel_type']}")
        
        return "; ".join(parts) if parts else "Configuration/System Type, Ahead Horse Power"
    
    def _format_owner_info(self, vessel, ai_details: Dict[str, Any], info_type: str) -> str:
        """Format owner or operator information"""
        name = ai_details.get(info_type) or getattr(vessel, info_type, None) or f"Line 1 = Official Name"
        address = ai_details.get(f'{info_type}_address') or getattr(vessel, f'{info_type}_location', None) or "Line 2 = City, State/Country"
        
        return f"{name}\n{address}"
    
    def _generate_complete_roi_from_evidence(self, ai_assistant) -> Dict[str, Any]:
        """Generate all ROI content directly from evidence using comprehensive AI analysis"""
        import logging
        logger = logging.getLogger('app')
        
        try:
            # Gather all evidence content
            evidence_content = self._gather_all_evidence_content()
            
            if not evidence_content:
                logger.error("🔴 DIRECT ROI AI: No evidence content available")
                return {}
            
            logger.info(f"🟡 DIRECT ROI AI: Analyzing {len(evidence_content)} characters of evidence")
            
            # Comprehensive prompt to extract ALL ROI information
            prompt = f"""
You are an expert USCG marine casualty investigator analyzing evidence to create a complete Report of Investigation.

EVIDENCE CONTENT:
{evidence_content[:50000] if len(evidence_content) > 50000 else evidence_content}

Extract ALL information needed for a complete USCG ROI document. Analyze every detail and create professional content for each section.

Generate a comprehensive JSON response with the following sections:

1. INCIDENT SUMMARY:
- Extract basic incident information (date, time, location, vessel(s), type of casualty)
- Identify what happened and when

2. EXECUTIVE SUMMARY:
- Scene setting paragraph (4-6 sentences describing the operational context and incident)
- Outcomes paragraph (4-6 sentences describing response, casualties, and consequences)
- Causal factors paragraph (4-5 sentences identifying the cause and contributing factors)

3. VESSEL INFORMATION:
- Extract ALL vessel details for Section 2
- Official name, numbers, specifications, ownership, equipment

4. PERSONNEL CASUALTIES:
- All people involved, their status, injuries, and details

5. FINDINGS OF FACT:
- Convert all evidence into numbered factual statements
- Organize chronologically and by topic
- 15-25 professional findings

6. ANALYSIS:
- Identify causal factors using USCG Swiss Cheese methodology
- Generate professional analysis for each factor

7. CONCLUSIONS:
- Determine initiating event and causal factors
- Address regulatory violations and recommendations

8. ACTIONS TAKEN:
- Post-incident actions by all parties

9. RECOMMENDATIONS:
- Safety and administrative recommendations

Return as comprehensive JSON:
{{
  "incident_summary": {{
    "date": "incident date",
    "time": "incident time", 
    "location": "specific location",
    "vessel_name": "primary vessel",
    "incident_type": "type of casualty",
    "description": "brief description"
  }},
  "executive_summary": {{
    "scene_setting": "4-6 sentence paragraph setting scene and describing incident",
    "outcomes": "4-6 sentence paragraph describing response and outcomes",
    "causal_factors": "4-5 sentence paragraph identifying cause and factors"
  }},
  "vessel_information": {{
    "official_name": "vessel name",
    "official_number": "O.N. number",
    "specifications": "comprehensive vessel details",
    "ownership": "owner and operator information",
    "equipment": "safety and operational equipment"
  }},
  "personnel_casualties": [
    {{
      "name": "person name",
      "role": "position",
      "status": "injured/deceased/uninjured",
      "details": "injury details and treatment"
    }}
  ],
  "findings_of_fact": [
    "4.1.1. Professional finding statement",
    "4.1.2. Second finding statement",
    "...continue with all findings"
  ],
  "causal_factors": [
    {{
      "title": "Failure of/Inadequate/Lack of [factor title]",
      "category": "organization/workplace/precondition/production/defense",
      "description": "detailed description",
      "analysis": "comprehensive analysis of how this factor contributed"
    }}
  ],
  "conclusions": {{
    "initiating_event": "The initiating event was...",
    "causal_determinations": ["Factor 1", "Factor 2", "Factor 3"],
    "violations": "Any regulatory violations identified",
    "other_conclusions": "Additional conclusions"
  }},
  "actions_taken": [
    "7.1. Action taken by Coast Guard/parties",
    "7.2. Additional actions taken",
    "...continue with all actions"
  ],
  "recommendations": {{
    "safety_recommendations": [
      "8.1.1. Specific safety recommendation",
      "8.1.2. Additional safety recommendation"
    ],
    "administrative_recommendations": [
      "8.2.1. Administrative recommendation",
      "...or empty array if none"
    ]
  }}
}}

CRITICAL REQUIREMENTS:
- Extract EVERYTHING mentioned in the evidence
- Make reasonable assumptions about standard maritime practices
- Generate 15-25 professional findings of fact
- Identify multiple causal factors (typically 3-7)
- Create specific, actionable recommendations
- Use professional USCG language and format
- Ensure all sections are comprehensive and complete

This ROI will be generated entirely from this evidence analysis, so be thorough and extract every relevant detail.
"""
            
            logger.info("🟡 DIRECT ROI AI: Sending comprehensive analysis request to AI")
            response = ai_assistant.chat(prompt)
            
            # Parse the comprehensive response
            roi_content = ai_assistant._safe_json_extract(response)
            
            if roi_content and isinstance(roi_content, dict):
                logger.info("🟢 DIRECT ROI AI: Successfully extracted comprehensive ROI content")
                logger.info(f"🟢 DIRECT ROI AI: Sections generated: {list(roi_content.keys())}")
                
                # Log content summaries
                if roi_content.get('findings_of_fact'):
                    logger.info(f"🟢 DIRECT ROI AI: Generated {len(roi_content['findings_of_fact'])} findings")
                if roi_content.get('causal_factors'):
                    logger.info(f"🟢 DIRECT ROI AI: Identified {len(roi_content['causal_factors'])} causal factors")
                if roi_content.get('actions_taken'):
                    logger.info(f"🟢 DIRECT ROI AI: Found {len(roi_content['actions_taken'])} actions taken")
                
                return roi_content
            else:
                logger.error("🔴 DIRECT ROI AI: Invalid response structure from AI")
                return {}
                
        except Exception as e:
            logger.error(f"🔴 DIRECT ROI AI: Error generating ROI content: {e}")
            return {}
    
    def _generate_ai_executive_summary(self, roi_content: Dict[str, Any]) -> None:
        """Generate executive summary using AI-extracted content"""
        if not self.document:
            return
        
        # Generate title from AI content
        title = self._generate_ai_title(roi_content)
        
        # Title paragraph
        title_para = self.document.add_paragraph()
        title_para.add_run(title).bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        self.document.add_paragraph()
        
        # Executive Summary heading
        summary_heading = self.document.add_paragraph()
        summary_heading.add_run("EXECUTIVE SUMMARY").bold = True
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        
        # Add AI-generated executive summary paragraphs
        exec_summary = roi_content.get('executive_summary', {})
        
        if exec_summary.get('scene_setting'):
            self.document.add_paragraph(exec_summary['scene_setting'])
            self.document.add_paragraph()
        
        if exec_summary.get('outcomes'):
            self.document.add_paragraph(exec_summary['outcomes'])
            self.document.add_paragraph()
        
        if exec_summary.get('causal_factors'):
            self.document.add_paragraph(exec_summary['causal_factors'])
        
        self.document.add_page_break()
    
    def _generate_ai_investigating_officers_report(self, roi_content: Dict[str, Any]) -> None:
        """Generate the investigating officer's report using AI-extracted content"""
        if not self.document:
            return
        
        # Header block
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Add date and control number
        p = self.document.add_paragraph()
        p.add_run("16732\n")
        p.add_run(self._format_date(datetime.now()))

        # Add spacing
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Title
        title = self._generate_ai_title(roi_content)
        title_para = self.document.add_paragraph()
        title_para.add_run(title).bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()

        # INVESTIGATING OFFICER'S REPORT heading
        report_heading = self.document.add_paragraph()
        report_heading.add_run("INVESTIGATING OFFICER'S REPORT").bold = True
        report_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()

        # Add AI section methods to this instance
        from src.models.roi_ai_sections import add_ai_section_methods
        add_ai_section_methods(self)
        
        # Generate all sections using AI content
        self._generate_ai_section_1_preliminary_statement(roi_content)
        self._generate_ai_section_2_vessels_involved(roi_content)
        self._generate_ai_section_3_personnel_casualties(roi_content)
        self._generate_ai_section_4_findings_of_fact(roi_content)
        self._generate_ai_section_5_analysis(roi_content)
        self._generate_ai_section_6_conclusions(roi_content)
        self._generate_ai_section_7_actions_taken(roi_content)
        self._generate_ai_section_8_recommendations(roi_content)

        # Add signature block
        self._generate_signature_block()
    
    def _generate_ai_title(self, roi_content: Dict[str, Any]) -> str:
        """Generate USCG-format title from AI content"""
        incident = roi_content.get('incident_summary', {})
        vessel_info = roi_content.get('vessel_information', {})
        
        # Vessel name and official number
        vessel_name = vessel_info.get('official_name', 'VESSEL')
        official_number = vessel_info.get('official_number', '')
        
        if official_number:
            vessel_text = f"{vessel_name.upper()} ({official_number})"
        else:
            vessel_text = vessel_name.upper()
        
        # Incident type with casualties
        incident_type = incident.get('incident_type', 'MARINE CASUALTY').upper()
        
        # Check for casualties in personnel section
        personnel = roi_content.get('personnel_casualties', [])
        has_fatalities = any(p.get('status', '').lower() in ['deceased', 'death'] for p in personnel)
        has_injuries = any(p.get('status', '').lower() in ['injured', 'injury'] for p in personnel)
        
        if has_fatalities:
            casualty_desc = f"{incident_type} WITH LOSS OF LIFE"
        elif has_injuries:
            casualty_desc = f"{incident_type} WITH INJURIES"
        else:
            casualty_desc = incident_type
        
        # Location
        location = incident.get('location', 'LOCATION').upper()
        if "near" not in location.lower() and "on" not in location.lower():
            location = f"ON {location}"
        
        # Date
        incident_date = incident.get('date', 'DATE')
        if incident_date != 'DATE':
            try:
                # Try to parse and format the date
                from datetime import datetime
                parsed_date = datetime.strptime(incident_date, '%Y-%m-%d')
                date_str = parsed_date.strftime('%B %d, %Y').upper()
            except:
                date_str = incident_date.upper()
        else:
            date_str = 'DATE'
        
        return f"{vessel_text}, {casualty_desc} {location} ON {date_str}"