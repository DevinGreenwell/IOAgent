# AI-generated ROI sections for direct evidence-to-ROI generation
# These methods generate ROI sections directly from AI-extracted content

from typing import Dict, Any, List
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ai_section_1_preliminary_statement(self, roi_content: Dict[str, Any]) -> None:
    """Section 1: Preliminary Statement using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("1. Preliminary Statement").bold = True

    # 1.1 - Authority statement
    self.document.add_paragraph(
        "1.1. This marine casualty investigation was conducted and this report was submitted in "
        "accordance with Title 46, Code of Federal Regulations (CFR), Subpart 4.07, and under the "
        "authority of Title 46, United States Code (USC) Chapter 63."
    )

    # 1.2 - Parties in interest from AI content
    vessel_info = roi_content.get('vessel_information', {})
    personnel = roi_content.get('personnel_casualties', [])
    
    parties = []
    if vessel_info.get('official_name'):
        parties.append(f"the owner of the {vessel_info['official_name']}")
    
    # Add key personnel as parties
    for person in personnel:
        if person.get('role', '').lower() in ['master', 'captain', 'operator']:
            parties.append(f"the {person['role'].lower()}")

    if parties:
        parties_text = ", ".join(parties)
        self.document.add_paragraph(
            f"1.2. The Investigating Officer designated {parties_text}, as parties-in-interest in this "
            f"investigation. No other individuals, organizations, or parties were designated a party-in-"
            f"interest in accordance with 46 CFR Subsection 4.03-10."
        )

    # 1.3 - Lead agency statement
    self.document.add_paragraph(
        "1.3. The Coast Guard was lead agency for all evidence collection activities involving this "
        "investigation."
    )

    # Check for fatalities from AI content
    has_fatality = any(p.get('status', '').lower() in ['deceased', 'death'] for p in personnel)

    if has_fatality:
        self.document.add_paragraph(
            "Due to this investigation involving a loss of life, the Coast Guard Investigative "
            "Service (CGIS) was notified and agreed to provide technical assistance as required."
        )

    self.document.add_paragraph(
        "No other persons or organizations assisted in this investigation."
    )

    # 1.4 - Time format statement
    incident_summary = roi_content.get('incident_summary', {})
    location = incident_summary.get('location', 'local time')
    self.document.add_paragraph(
        f"1.4. All times listed in this report are approximate and expressed in local time using a 24‑hour format."
    )

    self.document.add_paragraph()

def generate_ai_section_2_vessels_involved(self, roi_content: Dict[str, Any]) -> None:
    """Section 2: Vessels Involved using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("2. Vessels Involved in the Incident").bold = True

    vessel_info = roi_content.get('vessel_information', {})
    
    # Add photo placeholder
    self.document.add_paragraph()
    photo_para = self.document.add_paragraph()
    photo_para.add_run("Figure 1. Undated Photograph of Vessel").italic = True
    self.document.add_paragraph()

    # Create vessel information table
    table = self.document.add_table(rows=13, cols=2)
    table.style = 'Table Grid'

    # Populate with AI-extracted vessel information
    vessel_data = [
        ("Official Name:", vessel_info.get('official_name', 'Not documented')),
        ("", f"({vessel_info.get('official_name', 'VESSEL NAME').upper()})"),
        ("Identification Number:", vessel_info.get('official_number', 'Not documented')),
        ("Flag:", vessel_info.get('flag', 'United States')),
        ("Vessel Class/Type/Sub-Type", vessel_info.get('specifications', 'Not documented')),
        ("Build Year:", vessel_info.get('build_year', 'Not documented')),
        ("Gross Tonnage:", vessel_info.get('gross_tonnage', 'Not documented')),
        ("Length:", vessel_info.get('length', 'Not documented')),
        ("Beam/Width:", vessel_info.get('beam', 'Not documented')),
        ("Draft/Depth:", vessel_info.get('draft', 'Not documented')),
        ("Main/Primary Propulsion:", vessel_info.get('propulsion', 'Not documented')),
        ("Owner:", vessel_info.get('ownership', 'Not documented')),
        ("Operator:", vessel_info.get('operator', 'Not documented'))
    ]

    # Fill table
    for i, (label, value) in enumerate(vessel_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        # Make vessel name italic in second row
        if i == 1:
            for run in row.cells[1].paragraphs[0].runs:
                run.italic = True

    self.document.add_paragraph()

def generate_ai_section_3_personnel_casualties(self, roi_content: Dict[str, Any]) -> None:
    """Section 3: Personnel Casualties using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("3. Deceased, Missing, and/or Injured Persons").bold = True
    
    personnel = roi_content.get('personnel_casualties', [])
    casualties = [p for p in personnel if p.get('status', '').lower() in ['deceased', 'missing', 'injured']]
    
    if casualties:
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Light Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Relationship to Vessel"
        header_cells[1].text = "Sex"
        header_cells[2].text = "Age"
        header_cells[3].text = "Status"
        
        # Add casualty entries
        for person in casualties:
            row_cells = table.add_row().cells
            row_cells[0].text = person.get('role', 'Unknown')
            row_cells[1].text = person.get('sex', 'Unknown')
            row_cells[2].text = str(person.get('age', 'Unknown'))
            row_cells[3].text = person.get('status', 'Unknown').title()
    else:
        self.document.add_paragraph("No personnel casualties resulted from this incident.")
    
    self.document.add_paragraph()

def generate_ai_section_4_findings_of_fact(self, roi_content: Dict[str, Any]) -> None:
    """Section 4: Findings of Fact using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("4. Findings of Fact").bold = True
    
    # 4.1 The Incident
    subheading = self.document.add_paragraph()
    subheading.add_run("4.1. The Incident:").bold = True
    
    findings = roi_content.get('findings_of_fact', [])
    
    if findings:
        for finding in findings:
            self.document.add_paragraph(finding)
    else:
        self.document.add_paragraph("4.1.1. No specific findings of fact were documented.")
    
    self.document.add_paragraph()

def generate_ai_section_5_analysis(self, roi_content: Dict[str, Any]) -> None:
    """Section 5: Analysis using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("5. Analysis").bold = True
    
    intro_para = self.document.add_paragraph()
    intro_para.add_run("The Coast Guard's investigation identified the following causal factors that contributed to this marine casualty:")
    self.document.add_paragraph()
    
    causal_factors = roi_content.get('causal_factors', [])
    
    if causal_factors:
        for i, factor in enumerate(causal_factors, 1):
            # Analysis heading
            subheading = self.document.add_paragraph()
            subheading_run = subheading.add_run(f"5.{i}. {factor.get('title', 'Unknown Factor')}")
            subheading_run.bold = True
            
            # Analysis text
            analysis_para = self.document.add_paragraph()
            analysis_para.add_run(factor.get('analysis', factor.get('description', 'No analysis provided.')))
            
            self.document.add_paragraph()
    else:
        self.document.add_paragraph("No causal factors have been identified for analysis at this time.")
    
    self.document.add_paragraph()

def generate_ai_section_6_conclusions(self, roi_content: Dict[str, Any]) -> None:
    """Section 6: Conclusions using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("6. Conclusions").bold = True
    
    conclusions = roi_content.get('conclusions', {})
    
    # 6.1 Determination of Cause
    subheading = self.document.add_paragraph()
    subheading.add_run("6.1. Determination of Cause:").bold = True
    
    # Initiating event
    if conclusions.get('initiating_event'):
        self.document.add_paragraph(f"6.1.1. {conclusions['initiating_event']}")
    else:
        self.document.add_paragraph(
            "6.1.1. The initiating event for this casualty could not be conclusively determined."
        )
    
    # Causal determinations
    determinations = conclusions.get('causal_determinations', [])
    for i, determination in enumerate(determinations, 1):
        self.document.add_paragraph(f"6.1.1.{i}. {determination}")
    
    # Standard sections
    standard_sections = [
        "6.2. Evidence of Act(s) or Violation(s) of Law by Credentialed Mariners: None identified.",
        "6.3. Evidence of Act(s) or Violation(s) of Law by U.S. Coast Guard personnel or others: None identified.",
        "6.4. Evidence of Act(s) Subject to Civil Penalty: None identified.",
        "6.5. Evidence of Criminal Act(s): None identified.",
        "6.6. Need for New or Amended U.S. Law or Regulation: None identified."
    ]
    
    for section in standard_sections:
        self.document.add_paragraph(section)
    
    self.document.add_paragraph()

def generate_ai_section_7_actions_taken(self, roi_content: Dict[str, Any]) -> None:
    """Section 7: Actions Taken using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("7. Actions Taken Since the Incident").bold = True
    
    actions = roi_content.get('actions_taken', [])
    
    if actions:
        for action in actions:
            self.document.add_paragraph(action)
    else:
        self.document.add_paragraph(
            "7.1. The Coast Guard initiated a formal investigation under 46 CFR Part 4."
        )
    
    self.document.add_paragraph()

def generate_ai_section_8_recommendations(self, roi_content: Dict[str, Any]) -> None:
    """Section 8: Recommendations using AI content"""
    if not self.document:
        return
        
    heading = self.document.add_paragraph()
    heading.add_run("8. Recommendations").bold = True
    
    recommendations = roi_content.get('recommendations', {})
    
    # 8.1 Safety Recommendations
    subheading = self.document.add_paragraph()
    subheading.add_run("8.1. Safety Recommendations:").bold = True
    
    safety_recs = recommendations.get('safety_recommendations', [])
    if safety_recs:
        for rec in safety_recs:
            self.document.add_paragraph(rec)
    else:
        self.document.add_paragraph(
            "8.1.1. Conduct a comprehensive review of vessel safety procedures and emergency response protocols."
        )
    
    # 8.2 Administrative Recommendations
    self.document.add_paragraph()
    subheading = self.document.add_paragraph()
    subheading.add_run("8.2. Administrative Recommendations:").bold = True
    
    admin_recs = recommendations.get('administrative_recommendations', [])
    if admin_recs:
        for rec in admin_recs:
            self.document.add_paragraph(rec)
    else:
        self.document.add_paragraph("None at this time.")
    
    self.document.add_paragraph()

# Add these methods to the USCGROIGenerator class
def add_ai_section_methods(roi_generator):
    """Add AI section generation methods to ROI generator instance"""
    roi_generator._generate_ai_section_1_preliminary_statement = generate_ai_section_1_preliminary_statement.__get__(roi_generator)
    roi_generator._generate_ai_section_2_vessels_involved = generate_ai_section_2_vessels_involved.__get__(roi_generator)
    roi_generator._generate_ai_section_3_personnel_casualties = generate_ai_section_3_personnel_casualties.__get__(roi_generator)
    roi_generator._generate_ai_section_4_findings_of_fact = generate_ai_section_4_findings_of_fact.__get__(roi_generator)
    roi_generator._generate_ai_section_5_analysis = generate_ai_section_5_analysis.__get__(roi_generator)
    roi_generator._generate_ai_section_6_conclusions = generate_ai_section_6_conclusions.__get__(roi_generator)
    roi_generator._generate_ai_section_7_actions_taken = generate_ai_section_7_actions_taken.__get__(roi_generator)
    roi_generator._generate_ai_section_8_recommendations = generate_ai_section_8_recommendations.__get__(roi_generator)