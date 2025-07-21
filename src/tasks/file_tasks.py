"""File processing async tasks."""

import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import hashlib
import mimetypes

from src.celery_app import celery_app
from src.models.evidence import Evidence
from src.models.project import Project
from src.models.user import db
from src.utils.file_processor import FileProcessor

logger = logging.getLogger(__name__)


@celery_app.task(bind=True, max_retries=3)
def process_uploaded_file_async(self, evidence_id: int) -> Dict[str, Any]:
    """
    Process uploaded evidence file asynchronously.
    
    Args:
        evidence_id: ID of the evidence record
    
    Returns:
        Dictionary with processing results
    """
    try:
        logger.info(f"Starting file processing for evidence {evidence_id}")
        
        self.update_state(state='PROGRESS', meta={'status': 'Loading evidence record...'})
        
        with celery_app.app.app_context():
            evidence = Evidence.query.get(evidence_id)
            if not evidence:
                raise ValueError(f"Evidence {evidence_id} not found")
            
            file_path = Path(evidence.file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            self.update_state(state='PROGRESS', meta={'status': 'Processing file...'})
            
            # Initialize file processor
            processor = FileProcessor()
            
            # Extract text content based on file type
            if evidence.file_type == 'application/pdf':
                content = processor.extract_pdf_text(str(file_path))
            elif evidence.file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                content = processor.extract_docx_text(str(file_path))
            elif evidence.file_type.startswith('text/'):
                content = processor.extract_text_file(str(file_path))
            elif evidence.file_type.startswith('image/'):
                # For images, we might use OCR in the future
                content = "[Image file - content extraction not yet implemented]"
            else:
                content = f"[Unsupported file type: {evidence.file_type}]"
            
            # Calculate file hash for integrity
            file_hash = processor.calculate_file_hash(str(file_path))
            
            # Extract metadata
            metadata = processor.extract_file_metadata(str(file_path))
            
            # Update evidence record
            evidence.content = content
            evidence.file_hash = file_hash
            
            if not evidence.metadata:
                evidence.metadata = {}
            evidence.metadata.update({
                'file_info': metadata,
                'processed': True,
                'processed_at': datetime.utcnow().isoformat(),
                'content_length': len(content) if content else 0
            })
            
            # Generate preview if applicable
            if evidence.file_type.startswith('image/'):
                self.update_state(state='PROGRESS', meta={'status': 'Generating preview...'})
                preview_path = processor.generate_image_preview(str(file_path), evidence_id)
                evidence.metadata['preview_path'] = preview_path
            
            db.session.commit()
            
            logger.info(f"File processing completed for evidence {evidence_id}")
            
            # Trigger AI analysis if content was extracted
            if content and len(content) > 100:
                from src.tasks.ai_tasks import analyze_evidence_async
                analyze_evidence_async.delay(evidence_id)
            
            return {
                'status': 'success',
                'evidence_id': evidence_id,
                'content_extracted': bool(content),
                'content_length': len(content) if content else 0,
                'file_hash': file_hash,
                'message': 'File processed successfully'
            }
            
    except Exception as e:
        logger.error(f"Error processing file for evidence {evidence_id}: {str(e)}")
        retry_in = 2 ** self.request.retries
        raise self.retry(exc=e, countdown=retry_in)


@celery_app.task(bind=True)
def generate_file_preview_async(self, evidence_id: int) -> Dict[str, Any]:
    """
    Generate preview for evidence files.
    
    Args:
        evidence_id: ID of the evidence
    
    Returns:
        Dictionary with preview information
    """
    try:
        logger.info(f"Generating preview for evidence {evidence_id}")
        
        with celery_app.app.app_context():
            evidence = Evidence.query.get(evidence_id)
            if not evidence:
                raise ValueError(f"Evidence {evidence_id} not found")
            
            file_path = Path(evidence.file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            processor = FileProcessor()
            preview_info = {}
            
            if evidence.file_type == 'application/pdf':
                # Generate PDF preview (first page as image)
                preview_path = processor.generate_pdf_preview(str(file_path), evidence_id)
                preview_info['preview_path'] = preview_path
                preview_info['preview_type'] = 'image'
                
            elif evidence.file_type.startswith('image/'):
                # Generate thumbnail
                thumbnail_path = processor.generate_image_thumbnail(str(file_path), evidence_id)
                preview_info['thumbnail_path'] = thumbnail_path
                preview_info['preview_type'] = 'thumbnail'
                
            elif evidence.file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'text/plain']:
                # Generate text preview (first 500 characters)
                if evidence.content:
                    preview_info['text_preview'] = evidence.content[:500] + '...' if len(evidence.content) > 500 else evidence.content
                    preview_info['preview_type'] = 'text'
            
            # Update evidence metadata
            if not evidence.metadata:
                evidence.metadata = {}
            evidence.metadata['preview'] = preview_info
            db.session.commit()
            
            return {
                'status': 'success',
                'evidence_id': evidence_id,
                'preview_type': preview_info.get('preview_type'),
                'message': 'Preview generated successfully'
            }
            
    except Exception as e:
        logger.error(f"Error generating preview for evidence {evidence_id}: {str(e)}")
        raise


@celery_app.task(bind=True)
def batch_process_evidence_async(self, project_id: int) -> Dict[str, Any]:
    """
    Process all unprocessed evidence files for a project.
    
    Args:
        project_id: ID of the project
    
    Returns:
        Dictionary with batch processing results
    """
    try:
        logger.info(f"Starting batch evidence processing for project {project_id}")
        
        with celery_app.app.app_context():
            # Find unprocessed evidence
            unprocessed = Evidence.query.filter_by(project_id=project_id)\
                                       .filter(db.or_(
                                           Evidence.content == None,
                                           Evidence.content == ''
                                       )).all()
            
            if not unprocessed:
                return {
                    'status': 'success',
                    'message': 'No unprocessed evidence found',
                    'processed_count': 0
                }
            
            # Process each file
            processed = 0
            failed = 0
            
            for evidence in unprocessed:
                try:
                    process_uploaded_file_async.delay(evidence.id)
                    processed += 1
                except Exception as e:
                    logger.error(f"Failed to queue processing for evidence {evidence.id}: {str(e)}")
                    failed += 1
            
            return {
                'status': 'success',
                'message': f'Queued {processed} files for processing',
                'processed_count': processed,
                'failed_count': failed,
                'total_count': len(unprocessed)
            }
            
    except Exception as e:
        logger.error(f"Error in batch processing for project {project_id}: {str(e)}")
        raise


@celery_app.task(bind=True)
def validate_file_integrity_async(self, evidence_id: int) -> Dict[str, Any]:
    """
    Validate file integrity using hash comparison.
    
    Args:
        evidence_id: ID of the evidence
    
    Returns:
        Dictionary with validation results
    """
    try:
        logger.info(f"Validating file integrity for evidence {evidence_id}")
        
        with celery_app.app.app_context():
            evidence = Evidence.query.get(evidence_id)
            if not evidence:
                raise ValueError(f"Evidence {evidence_id} not found")
            
            file_path = Path(evidence.file_path)
            if not file_path.exists():
                return {
                    'status': 'error',
                    'evidence_id': evidence_id,
                    'valid': False,
                    'message': 'File not found'
                }
            
            processor = FileProcessor()
            current_hash = processor.calculate_file_hash(str(file_path))
            
            # Compare with stored hash
            stored_hash = evidence.file_hash or evidence.metadata.get('file_hash') if evidence.metadata else None
            
            if not stored_hash:
                # No hash stored, calculate and store it
                evidence.file_hash = current_hash
                db.session.commit()
                valid = True
                message = 'File hash calculated and stored'
            else:
                valid = current_hash == stored_hash
                message = 'File integrity verified' if valid else 'File integrity check failed - file may have been modified'
            
            # Update metadata
            if not evidence.metadata:
                evidence.metadata = {}
            evidence.metadata['last_integrity_check'] = {
                'timestamp': datetime.utcnow().isoformat(),
                'valid': valid,
                'hash': current_hash
            }
            db.session.commit()
            
            return {
                'status': 'success',
                'evidence_id': evidence_id,
                'valid': valid,
                'message': message,
                'hash': current_hash
            }
            
    except Exception as e:
        logger.error(f"Error validating file integrity for evidence {evidence_id}: {str(e)}")
        raise


class FileProcessor:
    """Helper class for file processing operations."""
    
    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return ""
    
    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    def extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file: {str(e)}")
                return ""
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def extract_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract file metadata."""
        import os
        from datetime import datetime
        
        stat = os.stat(file_path)
        
        return {
            'size_bytes': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'mime_type': mimetypes.guess_type(file_path)[0],
            'extension': Path(file_path).suffix.lower()
        }
    
    def generate_image_preview(self, file_path: str, evidence_id: int) -> str:
        """Generate preview for image files."""
        try:
            from PIL import Image
            
            # Create preview directory
            preview_dir = Path('static/previews')
            preview_dir.mkdir(exist_ok=True)
            
            # Generate preview
            img = Image.open(file_path)
            img.thumbnail((800, 600), Image.Resampling.LANCZOS)
            
            preview_filename = f"preview_{evidence_id}.jpg"
            preview_path = preview_dir / preview_filename
            img.save(preview_path, 'JPEG', quality=85)
            
            return f"/static/previews/{preview_filename}"
        except Exception as e:
            logger.error(f"Error generating image preview: {str(e)}")
            return None
    
    def generate_pdf_preview(self, file_path: str, evidence_id: int) -> Optional[str]:
        """Generate preview image from first page of PDF."""
        try:
            # This would require pdf2image library
            # For now, return None
            return None
        except Exception as e:
            logger.error(f"Error generating PDF preview: {str(e)}")
            return None
    
    def generate_image_thumbnail(self, file_path: str, evidence_id: int) -> Optional[str]:
        """Generate thumbnail for image files."""
        try:
            from PIL import Image
            
            # Create thumbnail directory
            thumb_dir = Path('static/thumbnails')
            thumb_dir.mkdir(exist_ok=True)
            
            # Generate thumbnail
            img = Image.open(file_path)
            img.thumbnail((200, 200), Image.Resampling.LANCZOS)
            
            thumb_filename = f"thumb_{evidence_id}.jpg"
            thumb_path = thumb_dir / thumb_filename
            img.save(thumb_path, 'JPEG', quality=85)
            
            return f"/static/thumbnails/{thumb_filename}"
        except Exception as e:
            logger.error(f"Error generating thumbnail: {str(e)}")
            return None


from datetime import datetime  # Add this import at the top of the file